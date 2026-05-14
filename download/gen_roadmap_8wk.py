#!/usr/bin/env python3
"""PARWA Production Readiness Roadmap - Merged 8-Week Plan
Combines original 6-week roadmap + code-verified cross-check audit items
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6); style.paragraph_format.line_spacing = 1.3

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hf = hs.font; hf.name = 'Calibri'; hf.bold = True
    if level == 1: hf.size = Pt(18); hf.color.rgb = RGBColor(0x0A, 0x16, 0x28)
    elif level == 2: hf.size = Pt(14); hf.color.rgb = RGBColor(0x0A, 0x16, 0x28)
    else: hf.size = Pt(12); hf.color.rgb = RGBColor(0x1A, 0x2B, 0x40)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]; cell.text = h
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F2027"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]; cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs: run.font.size = Pt(9)
        if r_idx % 2 == 0:
            for c_idx in range(len(headers)):
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F8FC"/>')
                row.cells[c_idx]._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths): row.cells[i].width = Inches(w)
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(10)
        r = p.add_run(text); r.font.size = Pt(10)
    else:
        r = p.add_run(text); r.font.size = Pt(10)
    return p

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    return p

def gate(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x0A, 0x82, 0x8A)
    return p

def tag(doc, label, text):
    """Add a paragraph with a colored tag prefix"""
    p = doc.add_paragraph()
    r = p.add_run(f'[{label}] '); r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B) if label in ('MISSING','BUG') else RGBColor(0x0A, 0x82, 0x8A)
    r = p.add_run(text); r.font.size = Pt(10)
    return p

# ══════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════
for _ in range(6): doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PARWA'); r.bold = True; r.font.size = Pt(42); r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Production Readiness Roadmap'); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Merged 8-Week Execution Plan'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x50, 0x60, 0x70)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Security Audit + Codebase Audit + Cross-Check Verification'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x60, 0x70, 0x80)

for _ in range(3): doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('32 items cross-checked against source code | 5 false positives removed | 23 confirmed missing'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Each week ends with mandatory testing gate before code push'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('May 2025'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x60, 0x70, 0x80)

doc.add_page_break()

# ══════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
for item in [
    '1. Executive Summary',
    '2. Cross-Check Audit Summary',
    '3. Week 1: Security Carry-Forward + HIGH Fixes + Quick Wins',
    '4. Week 2: CLARA RAG + FAKE Voting + Agent Lightning',
    '5. Week 3: AI Upgrades + Socket.io + Dashboard Pages',
    '6. Week 4: MEDIUM Security + LangGraph Nodes',
    '7. Week 5: RLS + Critical Bugs + Nginx + Infrastructure',
    '8. Week 6: Loopholes + JWT RS256 + Backup + MEDIUM Remaining',
    '9. Week 7: Comprehensive Testing',
    '10. Week 8: Production Deployment + Documentation',
    '11. Complete Gap-to-Week Mapping',
    '12. Git Workflow and Testing Protocol',
    '13. Effort Summary',
    '14. Success Metrics',
]:
    add_bullet(doc, item)
add_para(doc, '(Right-click the table of contents and select "Update Field" to refresh page numbers.)', size=9)
doc.add_page_break()

# ══════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

add_para(doc, 'This roadmap defines the path to take PARWA from ~82% built to 100% production-ready. It is a merged document combining three independent audit streams: (1) the original Production Readiness Security Audit identifying 93 findings across CRITICAL, HIGH, MEDIUM, and LOW severity levels; (2) a comprehensive codebase audit that cross-checked every claim against actual source code across 2,200+ Python files and ~145 frontend files; and (3) a code-verified cross-check audit of 32 claimed "genuinely missing" items.')

add_para(doc, 'The cross-check audit (Appendix B, full details) verified each of the 32 claimed missing items against the actual codebase. Results: 5 items were false positives (already fully implemented), 4 items were partial (mostly due to nginx config drift between two files), and 23 items are genuinely missing. This realistic assessment reduces the estimated effort and focuses development on verified gaps only.')

add_para(doc, 'Days 1-3 of the original roadmap (Auth Security, Data Protection, AI Rebuild) have been completed and cross-checked. 33 of 35 items verified as DONE, with 6 items carried forward to Week 1 as quick fixes. A critical finding: the security audit incorrectly claimed all 12 AI techniques were "pure regex" but cross-check proved all 12 use real LLM calls via llm_gateway.py with template fallbacks.')

add_para(doc, 'The remaining work is organized into 8 weekly phases. Weeks 1-6 cover the original roadmap items with false positives removed and cross-check evidence incorporated. Weeks 7-8 add comprehensive testing and production deployment validation that were under-served in the original 6-week plan. Each week concludes with a mandatory testing gate: unit tests, integration tests, and security regression. Code may only be pushed after all tests pass.')

add_para(doc, 'Current State: ~82% built. Backend AI engine is genuine (12 techniques with real LLM calls). Auth security hardened (JWT, httpOnly, MFA, roles). 2 of 7 dashboard pages fully built (Tickets 1,362 lines, Variants 722 lines). 5 dashboard pages are stubs. 6 critical bugs remain. 23 items genuinely missing. Nginx config drift between main and Docker configs is a critical discovery.')
add_para(doc, 'Target State: 100% production-ready. Zero CRITICAL/HIGH vulnerabilities. All DMD specifications implemented. CLARA RAG with HyDE + multi-query. MAKER with 6-24 LLM calls per query. FAKE Voting with consensus scoring. Agent Lightning self-learning. 25 Loophole Solutions. RLS on 93 tables. All 7 dashboards functional. Full test suite with 80%+ coverage. Production deployed with monitoring and documentation.')

# ══════════════════════════════════════════
# 2. CROSS-CHECK SUMMARY
# ══════════════════════════════════════════
doc.add_heading('2. Cross-Check Audit Summary', level=1)

add_para(doc, '32 items claimed as "genuinely missing" were verified against actual source code. Each item was checked by searching for relevant files, functions, and configurations. Full details in Appendix B.')

add_table(doc,
    ['Verdict', 'Count', 'Action', 'Impact'],
    [
        ['FALSE POSITIVE', '5', 'Remove from roadmap', 'Saves ~4-6 hours (UUID, .env, M-04, M-06, M-16 already built)'],
        ['PARTIAL', '4', 'Quick fixes', '2-3 hours (nginx config sync fixes 3 items + JWT blacklist)'],
        ['GENUINELY MISSING', '23', 'Build/Test', '90-105 hours across 8 weeks'],
    ],
    col_widths=[1.3, 0.6, 1.5, 3.6]
)

add_para(doc, 'Key Discovery: Nginx Config Drift. The file infra/docker/nginx-default.conf (runs in production via Docker) has significantly weaker security than nginx/nginx.conf. Main config has OCSP stapling, full HSTS with includeSubDomains/preload, and 8 cipher suites. Docker config has none of these. Syncing the Docker config resolves 3 partial items in one file change.')

add_para(doc, 'Days 1-3 Cross-Check: Day 1 (13/13 DONE), Day 2 (8/12 DONE, 4 carried forward), Day 3 (9/11 Met, 2 carried forward). Total carried forward: 6 items (C-07, C-15, C-14, H-09, AI-F01, AI-F02) assigned to Week 1 as quick fixes.')

# ══════════════════════════════════════════
# 3. WEEK 1
# ══════════════════════════════════════════
doc.add_heading('3. Week 1: Security Carry-Forward + HIGH Fixes + Quick Wins', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours | Dependencies: Days 1-3 completed')

doc.add_heading('3.1 Day 2-3 Carry-Forward Quick Fixes (2-3 hours)', level=2)
add_para(doc, 'These 6 items from Days 2-3 cross-check must be completed first before any larger work begins:')
for bp, txt in [
    ('[C-07] ', 'Remove .env.prod from git tracking. Run git rm --cached .env.prod, commit, add pre-commit hook preventing env files.'),
    ('[C-15] ', 'Add logging.warning() on startup when REFRESH_TOKEN_PEPPER is empty regardless of environment.'),
    ('[C-14] ', 'Create ADR documenting Fernet (AES-128-CBC + HMAC-SHA256) vs AES-256-GCM decision for OAuth tokens.'),
    ('[H-09] ', 'Add production ValueError guard for PRICING_SIGNING_KEY matching other config validators. Move from raw os.environ.get() in pricing.py line 663 to config.py.'),
    ('[AI-F01] ', 'Create backend/app/core/techniques/llm_client.py as thin re-export wrapper around llm_gateway.py (632 lines).'),
    ('[AI-F02] ', 'Add execute_with_llm() to BaseTechniqueNode in base.py centralizing LLM call patterns.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('3.2 Cross-Check Quick Wins (2-3 hours)', level=2)
add_para(doc, 'These are the quick infrastructure fixes identified in the cross-check audit that were not in the original roadmap:')
for bp, txt in [
    ('[CROSS-1] Nginx Config Sync: ', 'Copy security settings from nginx/nginx.conf to infra/docker/nginx-default.conf. This single change fixes OCSP stapling, HSTS includeSubDomains/preload, and expands TLS ciphers from 2 to 8. CRITICAL: Docker config is what runs in production.'),
    ('[CROSS-2] ENVIRONMENT Enum Validation: ', 'Add validator in config.py ensuring ENVIRONMENT is one of [development, staging, test, production]. Currently plain str accepts any value.'),
    ('[CROSS-3] Shared IP Extraction Utility: ', 'Create get_client_ip() in backend/app/core/utils.py. Currently 3 duplicate implementations: rate_limit.py line 121, ip_allowlist.py line 152, request_logger.py line 73.'),
    ('[CROSS-4] Nginx Non-Root: ', 'Add USER nginx directive to infra/docker/nginx.Dockerfile. chown exists at line 29 but container still runs as root.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('3.3 Security HIGH - Middleware (8-10 hours)', level=2)
for bp, txt in [
    ('[H-01] ', 'Open Redirect Fix: Validate redirect query parameter against whitelist. Block double-encoding. File: src/lib/auth-cookies.ts.'),
    ('[H-04] ', 'Content-Security-Policy: Add comprehensive CSP to security headers. File: backend/app/middleware/security_headers.py.'),
    ('[H-06] ', 'IP Extraction Standardization: Use shared get_client_ip() (built in 3.2) across rate_limit.py, request_logger.py, ip_allowlist.py.'),
    ('[H-13] ', 'Billing Role Restrictions: Add owner/admin checks to cancel/subscription/refund endpoints. File: backend/app/api/billing.py.'),
    ('[H-21] ', 'Auth Rate Limiting: Redis sliding window: Login 5/min/email, Register 3/min/IP, OTP 10/min/phone. File: backend/app/middleware/rate_limit.py.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('3.4 Security HIGH - Webhooks + CSRF + Remaining (8-10 hours)', level=2)
for bp, txt in [
    ('[H-07] ', 'Webhook Signature Verification: Fail-closed when webhook_secret missing. File: backend/app/api/billing_webhooks.py.'),
    ('[H-08] ', 'Webhook Replay Protection: Add timestamp validation rejecting requests older than 5 minutes. [CROSS-5 - verified missing, no timestamp check exists].'),
    ('[H-19] ', 'CSRF Protection: Server-side tokens in httpOnly cookies. Validate on POST/PUT/DELETE. File: backend/app/middleware/csrf.py.'),
    ('[H-20] ', 'Mock Login Production Gate: Disable when ENVIRONMENT=production. File: dashboard/src/app/api/auth/login/route.ts.'),
    ('[H-12] ', 'Google OAuth Token Exchange: Move token from URL query to POST body. File: backend/app/services/auth_service.py.'),
    ('[H-16] ', 'HTML Injection in Email Templates: Sanitize with bleach/markupsafe. File: dashboard/src/lib/notifications.ts.'),
    ('[H-17] ', 'Channel Status API Key Leakage: Remove sensitive chars from responses. Return only configured/not configured.'),
    ('[H-15] ', 'Webhook Management Auth: Add authentication to webhook query/retry endpoints. File: backend/app/api/webhooks.py.'),
    ('[M-27] ', 'User Enumeration Prevention: Identical response for existing/non-existing emails. File: src/app/api/auth/check-email/route.ts.'),
    ('[M-37] ', 'Hardcoded SMS Phone: Fix +1234567890 to actual customer phone. File: dashboard/src/lib/notifications.ts.'),
    ('[D-02/D-03] ', 'Analytics Mock Data + Email Content Sanitization: Connect to real APIs, route through backend.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('3.5 JWT Token Blacklist Implementation (2-3 hours)', level=2)
add_para(doc, '[CROSS-6 - PARTIAL] auth.py generates jti claim (line 98-99) and has get_token_jti() helper (line 149), but no is_token_revoked() function exists. Implement: Redis SET for blacklisted jtis with TTL matching token expiry, check on every authenticated request, Celery beat cleanup task for expired entries. This bridges to the full RS256 migration in Week 6.')

doc.add_heading('3.6 Week 1 Testing Gate', level=2)
for t in [
    'Unit tests for 6 carry-forward fixes (C-07, C-15, C-14, H-09, AI-F01, AI-F02)',
    'Unit tests for cross-check quick wins (nginx sync, ENV enum, shared IP, non-root)',
    'Integration tests for all HIGH security fixes (middleware, webhooks, CSRF, rate limiting)',
    'Unit tests for JWT blacklist (revoke, check, cleanup)',
    'Security regression: verify zero CRITICAL/HIGH introduced',
    'Coverage target: 80%+ on all modified files',
]: add_bullet(doc, t)
gate(doc, 'GATE: All tests pass -> git push -> proceed to Week 2')

doc.add_heading('3.7 Week 1 Acceptance Criteria', level=2)
add_para(doc, 'All 6 carry-forward items resolved. Nginx Docker config synced with main config (OCSP, HSTS, 8 ciphers). ENVIRONMENT enum validated. Shared IP utility in utils.py. Nginx runs as non-root. All 22 HIGH severity findings resolved. JWT blacklist functional. Rate limiting fails closed on Redis down. Webhook replay protection active. Zero CRITICAL vulnerabilities.')

# ══════════════════════════════════════════
# 4. WEEK 2
# ══════════════════════════════════════════
doc.add_heading('4. Week 2: CLARA RAG + FAKE Voting + Agent Lightning', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours | Dependencies: Week 1 complete (security baseline)')

doc.add_heading('4.1 CLARA RAG Advanced Retrieval (8-10 hours)', level=2)
add_para(doc, '[CROSS-7 - PARTIAL] CLARA exists as quality gate pipeline (clara_quality_gate.py, 704 lines) and RAG re-ranking (rag_reranking.py, 1,596 lines) but uses BM25/keyword scoring, NOT LLM. Context compression EXISTS (context_compression.py, 749 lines) with 5 strategies. Missing: HyDE and Multi-Query.')

add_para(doc, '[Build] HyDE (Hypothetical Document Embedding): Use LLM to generate hypothetical answer for customer query. Embed and use as primary retrieval vector. File: backend/app/core/rag/clara_rag.py (new).')
add_para(doc, '[Build] Multi-Query Retrieval: Generate 3 alternative query phrasings via LLM, retrieve for all 4 queries, merge/deduplicate, rank by aggregate score. File: backend/app/core/rag/multi_query.py (new).')
add_para(doc, '[Existing] Context Compression: 749 lines, 5 strategies (extractive, abstractive, hybrid, sliding_window, priority_based). No work needed.')
add_para(doc, '[Stretch] LLM Re-ranking: Upgrade rag_reranking.py from TF-IDF/BM25 to LLM-based relevance scoring (1-10 scale).')

doc.add_heading('4.2 FAKE Voting Sub-System (4-5 hours)', level=2)
add_para(doc, '[CROSS-8] MAKER validator (11_maker_validator.py, 684 lines) generates K solutions and scores but lacks FAKE Voting mechanism. Build: generate 3-5 candidates, multi-evaluator scoring (fluency, relevance, safety, brand voice), Red-Flagging for hallucinations/policy violations, weighted consensus scoring. Log results for Agent Lightning training data.')

doc.add_heading('4.3 Agent Lightning Training Pipeline (6-8 hours)', level=2)
add_para(doc, '[CROSS-9] training_tasks.py (140 lines) has 3 Celery tasks returning hardcoded zeros (samples_count=0, training_triggered=False). Database models exist (TrainingDataset, TrainingCheckpoint, AgentMistake, AgentPerformance, TrainingRun) but no logic connects them.')
add_para(doc, 'Build real pipeline: (1) Mistake collection from AgentMistake table for FAKE Voting red-flags + Jarvis overrides. (2) Data preparation converting to LLaMA-3-8B fine-tuning format via Unsloth. (3) Training trigger when mistakes exceed threshold (default 50). (4) Google Colab integration for training notebook. (5) Model deployment endpoint to upload and activate fine-tuned weights. (6) A/B testing framework splitting traffic between default and fine-tuned model.')

doc.add_heading('4.4 DSPy Integration + Smart Router Fixes (6-8 hours)', level=2)
add_para(doc, '[AI-12] Replace StubModule with real DSPy: Signature for customer care, Module chaining CoT + ReAct, Bayesian Signature Optimizer. File: backend/app/core/dspy/.')
add_para(doc, '[AI-16] 3-tier optimization: Tier 1 (Mini PARWA) fastest single technique, Tier 2 (PARWA) best single, Tier 3 (PARWA High) full MAKER.')
add_para(doc, '[BUG-HT] Smart Router HealthTracker: Refactor from Python dict to Redis-backed state for multi-worker consistency.')
add_para(doc, '[BUG-MV] MockVectorStore default: Reverse to use PgVectorStore when pgvector available. MockVectorStore only as fallback.')

doc.add_heading('4.5 Week 2 Testing Gate', level=2)
for t in [
    'Unit tests for CLARA RAG: HyDE generation, multi-query retrieval, dedup, ranking',
    'Unit tests for FAKE Voting: candidate generation, 4+ evaluator scoring, Red-Flagging, consensus',
    'Unit tests for Agent Lightning: mistake collection, data prep, threshold trigger, A/B framework',
    'Unit tests for DSPy: Signature validation, Module chaining, optimizer compilation',
    'Integration test: full query through CLARA RAG -> FAKE Voting with real LLM calls',
    'Integration test: Smart Router tier selection per variant',
    'Load test: 10 concurrent conversations through pipeline, document baselines',
    'Coverage target: 80%+ on all modified files',
]: add_bullet(doc, t)
gate(doc, 'GATE: All tests pass -> git push -> proceed to Week 3')

# ══════════════════════════════════════════
# 5. WEEK 3
# ══════════════════════════════════════════
doc.add_heading('5. Week 3: AI Upgrades + Socket.io + Dashboard Pages', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours | Dependencies: Week 2 complete (AI frameworks working)')

doc.add_heading('5.1 LLM-Powered AI Upgrades (10-12 hours)', level=2)

add_para(doc, '[CROSS-10] Voice Server LLM Intent: parwa_voice_server.py _detect_intent() (line 444) is pure keyword matching with 9 static categories. Connect to Smart Router for LLM-based classification. Keep keyword as fast path for obvious intents, use LLM for ambiguous ones. Comment at 447-448 says "in production routes through full pipeline" - make it true.')

add_para(doc, '[CROSS-11] LLM-Based Sentiment: sentiment_engine.py (855 lines) uses static word lists (FRUSTRATION_STRONG, POSITIVE_WORDS, EMOTION_LEXICON) and pattern matching (ALL CAPS, exclamation density). Add LLM layer producing structured output (sentiment_score: 0-1, emotion_labels: [], urgency_level: low/medium/high/critical). Keep lexicon as fast path, LLM for complex/ambiguous cases.')

add_para(doc, '[CROSS-12] LLM-Based PII NER: pii_redaction_engine.py (1,155 lines) uses 16 compiled regex patterns. Header says "no NLP libs". Add LLM-based NER as second pass for ambiguous cases where regex confidence is low (person names, addresses, custom entities). Hybrid: regex catches known patterns fast, LLM catches context-dependent PII.')

doc.add_heading('5.2 Socket.io Frontend Client (6-8 hours)', level=2)
add_para(doc, '[CROSS-13] Backend has full Socket.io server (socketio.py, 443 lines) with rooms, JWT auth, event buffering. Frontend has ZERO Socket.io code - socket.io-client not even in package.json.')
add_para(doc, 'Build: (1) Install socket.io-client. (2) Create useSocket React hook managing connection, auth token passing, reconnection with exponential backoff. (3) Real-time ticket status updates on Tickets page. (4) Live chat message streaming. (5) Notification push across all dashboard pages. (6) Event replay for missed messages during disconnection.')

doc.add_heading('5.3 Dashboard Pages (12-14 hours)', level=2)
add_para(doc, '[CROSS-14] 5 of 7 dashboard pages are stubs. Tickets (1,362 lines) and Variants (722 lines) are fully built. Build 5 pages with real backend API integration, Socket.io real-time updates, loading states, error handling, responsive design.')
add_para(doc, 'Pages at src/app/dashboard/: Billing (52 lines stub -> subscription, invoices, payments, upgrade/downgrade), Knowledge (52 lines stub -> upload, search, document list, reindex), Settings (15 lines stub -> profile, API keys, team, notifications, security tabs), AI Agents (15 lines stub -> agent list, metrics, capability matrix, training history), Monitoring (52 lines stub -> system health, AI metrics, alerts, SLA).')

doc.add_heading('5.4 Week 3 Testing Gate', level=2)
for t in [
    'Unit tests for sentiment LLM upgrade: structured output parsing, fallback behavior',
    'Unit tests for PII NER LLM: low-confidence regex triggers LLM, high-confidence skips',
    'Unit tests for voice intent: keyword fast path + LLM fallback + Smart Router integration',
    'Unit tests for useSocket hook: connect, auth, events, reconnect, missed replay',
    'Unit tests for all 5 dashboard pages: rendering, API calls, filters, error states',
    'Integration test: Socket.io connects to backend, receives real-time ticket updates',
    'Integration test: all 5 dashboard pages load real data from backend APIs',
    'Coverage target: 80%+ on all modified files',
]: add_bullet(doc, t)
gate(doc, 'GATE: All tests pass -> git push -> proceed to Week 4')

# ══════════════════════════════════════════
# 6. WEEK 4
# ══════════════════════════════════════════
doc.add_heading('6. Week 4: MEDIUM Security + LangGraph Nodes', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours | Dependencies: Week 3 complete')

doc.add_heading('6.1 Security MEDIUM Fixes (8-10 hours)', level=2)
add_para(doc, '14 genuinely missing MEDIUM items (M-04, M-06, M-16 confirmed as false positives in cross-check):')
for bp, txt in [
    ('[M-01] ', 'Remove user_role from AuthorizationError. Return generic "access denied". File: backend/app/api/deps.py.'),
    ('[M-05] ', 'Rate limiter fail-closed: block when Redis down. [CROSS-15 - rate_limit.py line 91 explicitly fail-opens].'),
    ('[M-08] ', 'Add auth dependency to events API. File: backend/app/main.py.'),
    ('[M-11] ', 'Cache-Control: no-store on auth responses. File: backend/app/middleware/security_headers.py.'),
    ('[M-13] ', 'Replace setattr-based user update with explicit field whitelist. File: backend/app/api/admin.py.'),
    ('[M-14] ', 'Add Pydantic models to ai_engine endpoints. File: backend/app/api/ai_engine.py.'),
    ('[M-17] ', 'Replace str(e) with generic errors in knowledge base API. File: backend/app/api/knowledge_base.py.'),
    ('[M-19] ', 'Fix visitor token pass-through on exception. File: backend/app/api/chat_widget.py.'),
    ('[M-23] ', 'Fix MCP server CORS ["*"] fallback. File: mcp_server/main.py.'),
    ('[M-26] ', 'Add security headers to Next.js middleware. File: src/middleware.ts.'),
    ('[M-28] ', 'Sanitize email content before Brevo API. File: dashboard/src/app/api/send-email/route.ts.'),
    ('[M-32] ', 'Max payload 1MB on Celery tasks. File: backend/app/tasks/celery_app.py.'),
    ('[M-33] ', 'Escape % and _ in ILIKE queries. Files: backend/app/api/admin.py, tickets.py.'),
    ('[M-35] ', 'Role check on notification endpoint. File: backend/app/api/notifications.py.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('6.2 LangGraph Node Completion (12-14 hours)', level=2)
add_para(doc, 'All 19 nodes must perform real processing. 12 need completion/rebuilding:')
for bp, txt in [
    ('[LG-01] Node 02 Empathy: ', 'LLM analyzes emotional state (anger/frustration/confusion/satisfaction/urgency), adjusts tone.'),
    ('[LG-02] Node 04 Base Domain: ', 'CLARA RAG retrieval (HyDE+Multi-Query) + Smart Router technique selection.'),
    ('[LG-03] Node 05 FAQ: ', 'Embedding similarity matching, LLM fallback below confidence threshold.'),
    ('[LG-04] Node 06 Refund: ', 'Connect to Paddle API for eligibility, amount calculation, initiation.'),
    ('[LG-05] Node 07 Technical: ', 'ReAct technique with diagnostic tools, KB search, step-by-step troubleshooting.'),
    ('[LG-06] Node 08 Billing: ', 'Subscription queries, invoice lookups, payment troubleshooting, plan changes.'),
    ('[LG-07] Node 09 Complaint: ', 'Sentiment analysis, service recovery playbooks, auto-escalation on severity.'),
    ('[LG-08] Node 10 Escalation: ', 'Determine escalation by topic/sentiment/confidence, route to human queue with context.'),
    ('[LG-09] Node 12 Control: ', 'Jarvis monitoring for accuracy, policy, tone, PII. Override/escalate below threshold.'),
    ('[LG-10] Node 13 DSPy: ', 'Connect to DSPy for technique selection optimization based on historical performance.'),
    ('[LG-11] Node 14 Guardrails: ', 'Block harmful content, enforce brand voice, check off-topic, ensure compliance. Integrate Loophole Engine here.'),
    ('[LG-12] Node 15 Channel Delivery: ', 'SMS (160 chars), HTML email, TTS voice, quick-reply chat per channel type.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('6.3 Week 4 Testing Gate', level=2)
for t in [
    'Unit tests for all 14 MEDIUM security fixes',
    'Unit tests for all 12 LangGraph nodes: input validation, LLM call structure, response parsing',
    'Integration test: full 19-node LangGraph pipeline with real LLM calls at each stage',
    'Integration test: all 7 dashboard pages with real data',
    'Multi-tenant isolation test: no cross-tenant data leakage',
    'Coverage target: 80%+ on all modified files',
]: add_bullet(doc, t)
gate(doc, 'GATE: All tests pass -> git push -> proceed to Week 5')

# ══════════════════════════════════════════
# 7. WEEK 5
# ══════════════════════════════════════════
doc.add_heading('7. Week 5: RLS + Critical Bugs + Infrastructure', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours | Dependencies: Week 4 complete')

doc.add_heading('7.1 Critical Database Bug Fix (2-3 hours)', level=2)
add_para(doc, '[CROSS-16 - FK Mismatch] Migrations 003 and 005 reference sessions.id FK but table renamed to tickets (BL01). No fix migration exists. Create migration 020 handling both fresh and existing databases. File: database/alembic/versions/020_fix_session_ticket_fk.py.')
add_para(doc, '[Note: BUG-2 UUID is FALSE POSITIVE] Cross-check confirmed both email_delivery_event.py and outbound_email.py use UUID(as_uuid=True) consistently. No migration needed.')

doc.add_heading('7.2 PostgreSQL Row-Level Security (10-12 hours)', level=2)
add_para(doc, '[CROSS-17] Zero RLS policies on any of 93 tenant-scoped tables. Build: (1) Alembic migration enabling RLS on all tables with company_id. (2) CREATE POLICY per table: restrict SELECT/UPDATE/INSERT/DELETE to matching tenant. (3) PostgreSQL function extracting company_id from JWT via SET app.current_tenant_id. (4) Modify connection factory to set tenant context on checkout. (5) RLS bypass role for migrations/admin. (6) Tests: direct SQL returns zero rows without context; admin works via superuser.')

doc.add_heading('7.3 Database Backup & Recovery (4-6 hours)', level=2)
add_para(doc, '[CROSS-18] Zero backup scripts exist. PostgreSQL runs in Docker with local volumes only. Build: (1) pg_dump script for daily full backups with timestamps. (2) WAL archiving for point-in-time recovery. (3) Automated cron in Docker. (4) Restore procedure tested on clean environment. (5) Disaster recovery runbook with RTO/RPO. (6) Automated backup verification on schedule.')

doc.add_heading('7.4 Infrastructure Hardening (6-8 hours)', level=2)
for bp, txt in [
    ('[L-01] ', 'JWT RS256 prep: Generate RSA-2048 key pair. (Full migration in Week 6.)'),
    ('[L-02] ', 'Token Blacklist cleanup: Celery beat task every 6 hours for expired entries (built Week 1, add cleanup).'),
    ('[L-04] ', 'Rate Limit Cleanup: Celery beat task removing expired Redis entries every 6 hours.'),
    ('[L-05] ', 'Circuit Breaker Thread Safety: threading.Lock for state transitions.'),
    ('[L-09] ', 'Async Login: Convert to async with async database session.'),
    ('[L-11] ', 'File Magic-Byte Validation: PNG (0x89504E47), PDF (0x25504446), JPEG (0xFFD8FF).'),
    ('[L-12] ', 'JWT Key Rotation: Multiple active keys with versioned metadata in Redis.'),
    ('[INF-01] ', 'Docker Security: non-root (nginx fixed Week 1), CPU/memory limits, read-only filesystem, health checks.'),
    ('[INF-02] ', 'Environment Variable Audit: .env.production.template (FALSE POSITIVE - .env.prod.example exists). Add startup validation.'),
    ('[M-24] ', 'Dev Port Binding: Restrict 5432/6379 to 127.0.0.1. File: docker-compose.yml.'),
    ('[M-25] ', 'Redis Dev Password: Add default password even in development.'),
    ('[M-38] ', 'Google AI Key: Move from URL query param to request header.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('7.5 AlertManager Diverse Receivers (1-2 hours)', level=2)
add_para(doc, '[CROSS-19] monitoring/alertmanager/alertmanager.yml has 2 identical webhook receivers to http://backend:8000/api/webhooks/alertmanager. If backend is down, all alerts are lost. Add email/Slack receiver as backup channel.')

doc.add_heading('7.6 Week 5 Testing Gate', level=2)
for t in [
    'Unit tests for FK migration: alembic upgrade head on fresh + existing databases',
    'Unit tests for RLS: direct SQL returns no rows without tenant, admin bypass works',
    'Integration test: multi-tenant isolation via both RLS and middleware (dual-layer)',
    'Integration test: backup script runs, restore succeeds, data integrity verified',
    'Infrastructure test: Docker non-root, resource limits, health checks pass',
    'Coverage target: 80%+ on all modified files',
]: add_bullet(doc, t)
gate(doc, 'GATE: All tests pass -> git push -> proceed to Week 6')

# ══════════════════════════════════════════
# 8. WEEK 6
# ══════════════════════════════════════════
doc.add_heading('8. Week 6: Loopholes + JWT RS256 + Remaining Items', level=1)
add_para(doc, 'Estimated Effort: 30-35 hours | Dependencies: Week 5 complete (RLS active, infrastructure hardened)')

doc.add_heading('8.1 25 Loophole Solutions Engine (10-12 hours)', level=2)
add_para(doc, '[CROSS-20] No centralized loophole detection exists. Loophole patterns in Building Codes doc and 15+ test files reference them, but no loophole_engine.py or LoopholeRegistry implementation. Build: (1) LoopholeRegistry dataclass mapping 25 categories to severity/detection/countermeasure. (2) LoopholeDetectionEngine in backend/app/core/loophole_engine.py with rule-based + LLM analysis. (3) Integration into LangGraph Node 14 Guardrails as mandatory pre-delivery check. Categories: hallucination, PII leakage, unauthorized access, emotional manipulation, biased responses, off-topic, escalation failures, brand voice, regulatory non-compliance, etc.')

doc.add_heading('8.2 JWT RS256 Migration (4-6 hours)', level=2)
add_para(doc, '[CROSS-21] Currently HS256 only (auth.py line 27). RSA keys generated in Week 5 prep. Implement: (1) Configure backend signing with RS256 private key. (2) Configure frontend verification with RS256 public key. (3) Key rotation support via Redis versioned metadata. (4) Graceful rotation without invalidating sessions. (5) Token blacklist (built Week 1) continues working with RS256.')

doc.add_heading('8.3 Incomplete Feature Completion (8-10 hours)', level=2)
for bp, txt in [
    ('[IC-03] ', 'Confidence Scoring Real Data: Wire formula to real sources (technique metrics, LLM quality, KB relevance, context).'),
    ('[IC-11] ', 'Hallucination Detection: Cross-check responses against KB facts. Verify numerical claims against real data. Flag unsupported statements.'),
    ('[IC-12] ', 'Self-Healing Engine: LLM API failure auto-retry with Smart Router fallback. DB connection auto-reconnect. Queue depth monitoring.'),
    ('[BUG-3] ', 'Route Prefix Mismatch: variant_check.py checks /api/v1/tickets but tickets.py uses /tickets. Align prefixes.'),
    ('[BUG-4] ', 'Jarvis Session Persistence: Migrate in-memory Map() to Redis-backed storage with TTL. JarvisSession models exist.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('8.4 Week 6 Testing Gate', level=2)
for t in [
    'Unit tests for all 25 loophole detection categories',
    'Unit tests for JWT RS256: signing, verification, key rotation without session invalidation',
    'Unit tests for hallucination detection, confidence scoring, self-healing',
    'Unit tests for route prefix alignment, Jarvis Redis persistence',
    'Integration test: RS256 token lifecycle (create, verify, revoke, rotate)',
    'Coverage target: 80%+ on all modified files',
]: add_bullet(doc, t)
gate(doc, 'GATE: All tests pass -> git push -> proceed to Week 7')

# ══════════════════════════════════════════
# 9. WEEK 7
# ══════════════════════════════════════════
doc.add_heading('9. Week 7: Comprehensive Testing', level=1)
add_para(doc, 'Estimated Effort: 30-35 hours | Dependencies: All development complete (Weeks 1-6)')
add_para(doc, 'Week 7 is dedicated entirely to testing. Every feature built in Weeks 1-6 is verified through unit tests, integration tests, load tests, failover tests, and security regression tests. This is the quality assurance phase that proves PARWA is production-ready.')

doc.add_heading('9.1 Full Regression Test Suite (8-10 hours)', level=2)
for bp, txt in [
    ('[TEST-01] ', 'Run pytest backend/app/tests/ with 80%+ coverage on all modified files. Fix any failing tests discovered.'),
    ('[TEST-02] ', 'Phase verification tests: targeted test per fix in Weeks 1-6. Each fix has verify + regression test.'),
    ('[TEST-03] ', 'AI technique integration: verify all 12 techniques make real LLM calls. Mock responses, verify call structure, prompt format, response parsing.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('9.2 End-to-End Pipeline Tests (6-8 hours)', level=2)
for bp, txt in [
    ('[TEST-04] ', 'Full 19-node LangGraph pipeline with real LLM calls. Test all 4 verticals (refund, technical, billing, complaint) x 3 tiers (Mini, PARWA, PARWA High) = 12 scenarios.'),
    ('[TEST-05] ', 'Customer journey: signup -> create ticket -> AI response -> escalation -> resolution. Full E2E including Socket.io real-time updates.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('9.3 Load Testing (4-6 hours)', level=2)
add_para(doc, '[TEST-06] 100 concurrent conversations processed simultaneously. Verify no connection pool exhaustion, no memory leaks, no rate limit errors, consistent response times. Document performance baselines for: API P50/P95 latency, LLM tokens per request, DB query times, Redis response times.')

doc.add_heading('9.4 Failover Testing (4-6 hours)', level=2)
for bp, txt in [
    ('[TEST-07a] ', 'Kill Redis -> verify fail-closed rate limiting (blocks requests, not allows all)'),
    ('[TEST-07b] ', 'Kill PostgreSQL -> verify graceful error handling and automatic reconnection'),
    ('[TEST-07c] ', 'Simulate LLM API failures -> verify Smart Router fallback to alternative providers'),
    ('[TEST-07d] ', 'Simulate webhook failures -> verify retry logic and dead letter queue'),
    ('[TEST-07e] ', 'Kill Celery workers -> verify tasks requeue, no data loss'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('9.5 Security + Multi-Tenant Tests (4-6 hours)', level=2)
for bp, txt in [
    ('[TEST-08] ', 'Security regression: verify all 93 findings (1 CRITICAL, 22 HIGH, 39 MEDIUM, 31 LOW) resolved. Create new tests for previously-CRITICAL/HIGH items.'),
    ('[TEST-09] ', 'Multi-tenant isolation: attempt cross-tenant data access from Company A to Company B via every API endpoint. Verify blocked by BOTH middleware AND RLS (dual-layer).'),
    ('[TEST-10] ', 'SSL/TLS: verify nginx Docker config has OCSP, HSTS includeSubDomains/preload, 8 ciphers. Run testssl.sh or SSL Labs.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('9.6 Week 7 Testing Gate', level=2)
for t in [
    'Full pytest suite passes with 80%+ overall coverage',
    '12 E2E pipeline scenarios pass (4 verticals x 3 tiers)',
    '100 concurrent conversations: documented baselines, no degradation',
    'All 5 failover scenarios recover correctly',
    'Multi-tenant isolation: zero cross-tenant access via any endpoint',
    'Security regression: zero CRITICAL/HIGH/MEDIUM findings remain',
    'SSL/TLS config verified: OCSP, HSTS, 8 ciphers in Docker config',
]: add_bullet(doc, t)
gate(doc, 'GATE: All tests pass -> git push -> proceed to Week 8')

# ══════════════════════════════════════════
# 10. WEEK 8
# ══════════════════════════════════════════
doc.add_heading('10. Week 8: Production Deployment + Documentation', level=1)
add_para(doc, 'Estimated Effort: 25-30 hours | Dependencies: Week 7 complete (all tests passing)')

doc.add_heading('10.1 Production Deployment (12-14 hours)', level=2)
for bp, txt in [
    ('[DEPLOY-01] ', 'Verify docker-compose.prod.yml: all services non-root (nginx fixed Week 1), resource limits, health checks, startup order.'),
    ('[DEPLOY-02] ', 'SSL/TLS: TLSv1.2+, strong ciphers (8 suites from nginx sync), OCSP stapling, HSTS with preload. Must pass SSL Labs A+.'),
    ('[DEPLOY-03] ', 'Database migration on production data: alembic upgrade head on existing DB. Test rollback. Verify zero data loss.'),
    ('[DEPLOY-04] ', 'Backup verification: restore from backup created in Week 5 on staging. Verify data integrity.'),
    ('[DEPLOY-05] ', 'Monitoring: Grafana dashboards + Prometheus + AlertManager (with diverse receivers from Week 5). Alerts: API error >1%, P95 >2s, queue >100, LLM error >5%.'),
    ('[DEPLOY-06] ', 'AI Pipeline Metrics: per-technique latency, token usage by tier, LLM error rates by provider, Smart Router fallback frequency, FAKE Voting consensus, Jarvis override rates.'),
    ('[DEPLOY-07] ', 'Error Tracking: Sentry with structured logging, request tracing, correlation IDs.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('10.2 Documentation Update (6-8 hours)', level=2)
for bp, txt in [
    ('[DOC-01] ', 'Architecture Document: Update to reflect real AI architecture. Remove unimplemented claims. Document LLM integration via llm_gateway.py, CLARA RAG (HyDE+Multi-Query), MAKER pipeline, FAKE Voting.'),
    ('[DOC-02] ', 'Document Audit: Audit all docs in /documents/ and /docs/. Fix incorrect claims. Ensure LLM-first with template fallback is accurately described.'),
    ('[DOC-03] ', 'Production Runbook: env setup, DB migration, Docker deploy, SSL, monitoring, rollback, troubleshooting.'),
    ('[DOC-04] ', 'API Documentation: Regenerate OpenAPI/Swagger for all auth changes, new endpoints, modified schemas.'),
    ('[DOC-05] ', 'Spec Documents: Update PARWA_AI_Technique_Framework.md, PARWA_Context_Bible.md, JARVIS_SPECIFICATION.md, PARWA_SRS.md.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('10.3 Week 8 Final Testing Gate', level=2)
for t in [
    'Production smoke test on staging',
    'SSL Labs A+ test (nginx Docker config now has full ciphers + OCSP + HSTS)',
    'Database migration on production data copy (zero data loss)',
    'Backup restore verification (data integrity confirmed)',
    'Monitoring alerts validation (trigger each condition, verify notification)',
    'Full E2E production simulation: signup -> ticket -> AI response -> escalation -> resolution',
    'Performance regression: compare against Week 7 baselines, no degradation',
    'Documentation accuracy: all docs match actual system behavior',
]: add_bullet(doc, t)
gate(doc, 'GATE: All tests pass -> PRODUCTION DEPLOY APPROVED')

# ══════════════════════════════════════════
# 11. GAP MAPPING
# ══════════════════════════════════════════
doc.add_heading('11. Complete Gap-to-Week Mapping', level=1)
add_para(doc, 'Every verified gap mapped to its week, with cross-check verdict (FP = false positive, PARTIAL, MISSING).')

add_table(doc,
    ['Week', 'ID', 'Category', 'Item', 'Verdict'],
    [
        ['W1', 'C-07/C-15/C-14/H-09', 'Security', 'Carry-forward fixes from Day 2-3', 'Missing'],
        ['W1', 'AI-F01/AI-F02', 'AI', 'LLM client wrapper + execute_with_llm()', 'Missing'],
        ['W1', 'CROSS-1', 'Infra', 'Nginx config sync (OCSP+HSTS+ciphers)', 'Partial->Fixed'],
        ['W1', 'CROSS-2', 'Security', 'ENVIRONMENT enum validation', 'Missing'],
        ['W1', 'CROSS-3', 'Code', 'Shared IP extraction utility', 'Missing'],
        ['W1', 'CROSS-4', 'Infra', 'Nginx non-root USER directive', 'Missing'],
        ['W1', 'H-01 to H-21', 'Security HIGH', '22 HIGH severity findings', 'Missing'],
        ['W1', 'CROSS-5', 'Security', 'Webhook replay timestamp validation', 'Missing'],
        ['W1', 'CROSS-6', 'Security', 'JWT token blacklist is_token_revoked()', 'Partial->Fixed'],
        ['W2', 'CROSS-7', 'AI RAG', 'CLARA RAG: HyDE + Multi-Query', 'Partial->Fixed'],
        ['W2', 'CROSS-8', 'AI MAKER', 'FAKE Voting multi-evaluator consensus', 'Missing'],
        ['W2', 'CROSS-9', 'AI Training', 'Agent Lightning real training pipeline', 'Missing'],
        ['W2', 'AI-12', 'AI DSPy', 'Real DSPy integration (no StubModule)', 'Missing'],
        ['W2', 'AI-16', 'AI Routing', '3-Tier Hybrid optimization', 'Missing'],
        ['W2', 'BUG-HT/MV', 'Bug', 'HealthTracker Redis + MockVectorStore swap', 'Missing'],
        ['W3', 'CROSS-10', 'AI Voice', 'LLM-based intent classification', 'Missing'],
        ['W3', 'CROSS-11', 'AI NLP', 'LLM-based sentiment analysis', 'Missing'],
        ['W3', 'CROSS-12', 'AI NER', 'LLM-based PII NER detection', 'Missing'],
        ['W3', 'CROSS-13', 'Frontend', 'Socket.io frontend client', 'Missing'],
        ['W3', 'CROSS-14', 'Frontend', '5 Dashboard pages (Billing/KB/Settings/Agents/Monitor)', 'Missing'],
        ['W4', 'M-01-35', 'Security MED', '14 MEDIUM fixes (3 FP removed)', 'Missing'],
        ['W4', 'LG-01-12', 'LangGraph', '12 LangGraph node implementations', 'Missing'],
        ['W5', 'CROSS-16', 'Database', 'FK mismatch migration fix', 'Missing'],
        ['W5', 'CROSS-17', 'Security', 'PostgreSQL RLS (93 tables)', 'Missing'],
        ['W5', 'CROSS-18', 'Infra', 'Database backup & recovery', 'Missing'],
        ['W5', 'L-01-12', 'Infra', 'Infrastructure hardening (12 items)', 'Missing'],
        ['W5', 'CROSS-19', 'Infra', 'AlertManager diverse receivers', 'Missing'],
        ['W6', 'CROSS-20', 'Safety', '25 Loophole Solutions engine', 'Missing'],
        ['W6', 'CROSS-21', 'Security', 'JWT RS256 migration', 'Missing'],
        ['W6', 'IC-03/11/12', 'Incomplete', 'Confidence/Hallucination/Self-Healing', 'Missing'],
        ['W6', 'BUG-3/4', 'Bug', 'Route prefix + Jarvis persistence', 'Missing'],
        ['W7', '-', 'Testing', 'Full regression + load + failover + security', 'New'],
        ['W8', '-', 'Deploy', 'Production deployment + monitoring + docs', 'New'],
        ['-', 'BUG-2 (UUID)', 'Database', 'UUID type inconsistency', 'FALSE POSITIVE'],
        ['-', '.env template', 'Infra', '.env.production.template', 'FALSE POSITIVE'],
        ['-', 'M-04/M-06/M-16', 'Security', 'email-validator/API key 401/Twilio sig', 'FALSE POSITIVE'],
    ],
    col_widths=[0.4, 1.2, 0.7, 3.2, 0.8]
)

# ══════════════════════════════════════════
# 12. GIT WORKFLOW
# ══════════════════════════════════════════
doc.add_heading('12. Git Workflow and Testing Protocol', level=1)

doc.add_heading('12.1 Strict Git Workflow', level=2)
for item in [
    'At start of every session: git pull origin main',
    'Create feature branch: git checkout -b week-N-task-name',
    'Complete assigned work for the week',
    'Run full testing gate for the week',
    'Fix ALL failing tests',
    'Only after ALL tests pass: git push origin week-N-task-name',
    'Create PR, merge to main, pull before next week',
]: add_bullet(doc, item)

doc.add_heading('12.2 Testing Protocol Per Week', level=2)
for bp, txt in [
    ('1. Unit Tests: ', 'pytest for all new/modified code. 80%+ coverage. pytest-cov reporting.'),
    ('2. Integration Tests: ', 'API endpoints, database ops, cross-service communication.'),
    ('3. Security Regression: ', 'No new CRITICAL/HIGH vulnerabilities.'),
    ('4. Performance Baseline: ', 'No degradation from previous week.'),
    ('5. Documentation: ', 'Docstrings on new code. OpenAPI/Swagger updated.'),
]: add_bullet(doc, txt, bold_prefix=bp)

doc.add_heading('12.3 Parallel Work Handling', level=2)
for item in [
    'Always pull latest main before starting any new work',
    'Resolve merge conflicts immediately',
    'Run full test suite after any merge',
    'Never push without passing all tests',
    'Rebase frequently against main when on multiple branches',
]: add_bullet(doc, item)

# ══════════════════════════════════════════
# 13. EFFORT SUMMARY
# ══════════════════════════════════════════
doc.add_heading('13. Effort Summary', level=1)

add_table(doc,
    ['Week', 'Focus Area', 'Hours', 'Primary Risk'],
    [
        ['Week 1', 'Carry-Forward + HIGH + Cross-Check Quick Wins', '35-40', 'FAKE Voting prompt quality; nginx sync'],
        ['Week 2', 'CLARA RAG + FAKE Voting + Agent Lightning + DSPy', '35-40', 'HyDE prompt engineering; training data format'],
        ['Week 3', 'LLM Upgrades (Sentiment/PII/Voice) + Socket.io + 5 Dashboards', '35-40', 'Dashboard API integration; Socket.io auth'],
        ['Week 4', '14 MEDIUM fixes + 12 LangGraph nodes', '35-40', 'Node interdependencies; MEDIUM fix scope'],
        ['Week 5', 'RLS + FK Bug + Nginx + Backups + Infrastructure', '35-40', 'RLS migration performance; backup verification'],
        ['Week 6', '25 Loopholes + JWT RS256 + Incomplete features', '30-35', 'Loophole detection accuracy; key rotation'],
        ['Week 7', 'Full Testing (regression + load + failover + security)', '30-35', 'Test coverage gaps; load test infrastructure'],
        ['Week 8', 'Production Deploy + Monitoring + Documentation', '25-30', 'SSL config; doc accuracy; staging issues'],
        ['TOTAL', '8 Weeks', '265-305', 'Realistic with parallel work: 4-5 weeks focused'],
    ],
    col_widths=[0.7, 3.0, 0.7, 2.6]
)

add_para(doc, 'The 8-week timeline includes significant buffer. Realistic focused effort is ~135-155 hours (4-5 weeks for a single full-time developer). With parallel work and the developer working on multiple items simultaneously, 8 weeks provides adequate room for iteration, unexpected issues, and production validation. Week 7 (testing) and Week 8 (deployment) were not adequately covered in the original 6-week plan and are the primary additions.')

# ══════════════════════════════════════════
# 14. SUCCESS METRICS
# ══════════════════════════════════════════
doc.add_heading('14. Success Metrics', level=1)
add_para(doc, 'The roadmap succeeds when these metrics are achieved at the end of Week 8:')

for m in [
    'Zero CRITICAL vulnerabilities (C-07 resolved Week 1)',
    'Zero HIGH vulnerabilities (all 22 resolved Week 1)',
    'All 14 MEDIUM findings resolved (Week 4, 3 false positives removed)',
    'All 12 AI techniques making real LLM calls (verified Day 3 cross-check)',
    'CLARA RAG with HyDE + multi-query (Week 2)',
    'FAKE Voting multi-evaluator consensus (Week 2)',
    'Agent Lightning real training pipeline (Week 2)',
    'LLM-powered sentiment/PII/voice (Week 3)',
    'Socket.io real-time across dashboard (Week 3)',
    'All 7 dashboards with real data (Week 3)',
    'All 19 LangGraph nodes real processing (Week 4)',
    'RLS on 93 tenant-scoped tables (Week 5)',
    'Nginx Docker config: OCSP + HSTS + 8 ciphers (Week 1)',
    'Database backup with verified restore (Week 5)',
    '25 Loophole Solutions active (Week 6)',
    'JWT RS256 with key rotation (Week 6)',
    'Full test suite: 80%+ coverage, no regressions (Week 7)',
    '100 concurrent conversations stable (Week 7)',
    'SSL Labs A+ (Week 8)',
    'Production deployed with monitoring (Week 8)',
    'All docs match reality (Week 8)',
]: add_bullet(doc, m)

# ══════════════════════════════════════════
# APPENDIX B
# ══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Appendix B: Full Cross-Check Audit Details', level=1)
add_para(doc, 'This appendix contains the complete evidence for every item verified during the cross-check audit. Each item was checked by reading the actual source file referenced in the claim.')

doc.add_heading('B.1 False Positives (5 items - Already Built)', level=2)
add_table(doc,
    ['Item', 'Claim', 'Evidence (File:Line)', 'Verdict'],
    [
        ['BUG-2: UUID', '2 models use UUID vs String(36)', 'email_delivery_event.py + outbound_email.py: ALL columns use UUID(as_uuid=True) consistently', 'FALSE POSITIVE'],
        ['.env template', 'No production template', '.env.prod.example (85 lines) + .env.example (99 lines) exist at root', 'FALSE POSITIVE'],
        ['M-04', 'Simple @ validation', 'requirements.txt:34 email-validator>=2.1.0. Imported in verification.py', 'FALSE POSITIVE'],
        ['M-06', 'Pass-through on missing header', 'api_key_auth.py:77-86 returns 401 for invalid API keys', 'FALSE POSITIVE'],
        ['M-16', 'No Twilio signature verify', 'sms_channel.py:747-798 full HMAC verification via verify_twilio_signature()', 'FALSE POSITIVE'],
    ],
    col_widths=[0.9, 1.3, 3.0, 0.9]
)

doc.add_heading('B.2 Partial Items (4 items - Quick Fixes)', level=2)
add_table(doc,
    ['Item', 'Exists', 'Missing', 'Evidence'],
    [
        ['OCSP Stapling', 'nginx/nginx.conf:100-101', 'NOT in infra/docker/nginx-default.conf', 'Docker config runs in prod but missing OCSP'],
        ['HSTS subdomains', 'nginx/nginx.conf:106 has includeSubDomains; preload', 'docker config:46 only max-age', 'Docker config weaker'],
        ['TLS Ciphers', 'nginx/nginx.conf: 8 ciphers', 'docker config:42 only 2 ciphers', 'Older browsers may fail TLS in prod'],
        ['JWT Blacklist', 'jti generated (auth.py:98), get_token_jti() (auth.py:149)', 'No is_token_revoked() anywhere', 'Documented but not implemented'],
    ],
    col_widths=[1.0, 1.7, 1.5, 1.8]
)

doc.add_heading('B.3 Dashboard Pages Verified (5 stubs, 2 real)', level=2)
add_table(doc,
    ['Page', 'Status', 'Lines', 'Path'],
    [
        ['Tickets', 'REAL', '1,362', 'src/app/dashboard/tickets/page.tsx'],
        ['Variants', 'REAL', '722', 'src/app/dashboard/variants/page.tsx'],
        ['Billing', 'STUB', '52', 'src/app/dashboard/billing/page.tsx'],
        ['Knowledge', 'STUB', '52', 'src/app/dashboard/knowledge/page.tsx'],
        ['Settings', 'STUB', '15', 'src/app/dashboard/settings/page.tsx'],
        ['Agents', 'STUB', '15', 'src/app/dashboard/agents/page.tsx'],
        ['Monitoring', 'STUB', '52', 'src/app/dashboard/monitoring/page.tsx'],
    ],
    col_widths=[1.0, 0.6, 0.5, 3.9]
)

# ── Save ──
output_path = '/home/z/my-project/download/PARWA_Production_Readiness_Roadmap.docx'
doc.save(output_path)
import os; size = os.path.getsize(output_path)
print(f'Saved: {output_path}')
print(f'Size: {size:,} bytes ({size/1024:.1f} KB)')
print(f'Paragraphs: {len(doc.paragraphs)} | Tables: {len(doc.tables)}')
