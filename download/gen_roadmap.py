#!/usr/bin/env python3
"""Generate PARWA Production Readiness Roadmap - Week-Wise Phases with Testing Gates"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.3

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hf = hs.font
    hf.name = 'Calibri'
    hf.bold = True
    if level == 1:
        hf.size = Pt(18)
        hf.color.rgb = RGBColor(0x0A, 0x16, 0x28)
    elif level == 2:
        hf.size = Pt(14)
        hf.color.rgb = RGBColor(0x0A, 0x16, 0x28)
    else:
        hf.size = Pt(12)
        hf.color.rgb = RGBColor(0x1A, 0x2B, 0x40)

# ── Helpers ──
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F2027"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    # Data
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        if r_idx % 2 == 0:
            for c_idx in range(len(headers)):
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F8FC"/>')
                row.cells[c_idx]._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r = p.add_run(text)
        r.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
    return p

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    return p

# ══════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PARWA')
r.bold = True
r.font.size = Pt(42)
r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Production Readiness Roadmap')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Week-Wise Execution Plan \u2014 6 Phases to Production')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x50, 0x60, 0x70)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Security Audit + Codebase Audit Merged  |  ~75% Built \u2192 100% Production Ready')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x60, 0x70, 0x80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('May 2025')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x60, 0x70, 0x80)

doc.add_page_break()

# ══════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Days 1-3 Cross-Check Verification',
    '3. Phase 1: Security Carry-Forward + Security HIGH (Week 1)',
    '4. Phase 2: AI Frameworks + Dashboard Batch 1 (Week 2)',
    '5. Phase 3: Security MEDIUM + Dashboard + LangGraph (Week 3)',
    '6. Phase 4: RLS + Critical Bugs + Infrastructure (Week 4)',
    '7. Phase 5: Loopholes + Incomplete Features + Testing (Week 5)',
    '8. Phase 6: Deployment Readiness + Documentation (Week 6)',
    '9. Complete Gap-to-Phase Mapping',
    '10. Git Workflow and Testing Protocol',
    '11. Effort Summary',
    '12. Success Metrics',
]
for item in toc_items:
    add_bullet(doc, item)

add_para(doc, '(Right-click on the table of contents above and select "Update Field" to refresh page numbers after opening in Word.)', size=9)

doc.add_page_break()

# ══════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

add_para(doc, 'This roadmap defines the precise path to take PARWA from its current state (~75% built) to full production readiness. It is a merged document combining two independent audit streams: (1) the original Production Readiness Security Audit identifying 93 findings across CRITICAL, HIGH, MEDIUM, and LOW severity levels, and (2) a comprehensive codebase audit that cross-checked every claim against actual source code across 2,200+ Python files, 16MB of backend code, and ~145 frontend files.')

add_para(doc, 'Days 1-3 of the original roadmap (Auth Security, Data Protection, AI Rebuild) have been completed by the developer and independently cross-checked against the actual codebase. The cross-check verified 33 out of 35 items as DONE, with 6 items requiring carry-forward to Phase 1. A critical finding during cross-check: the security audit incorrectly claimed all 12 AI techniques were "pure regex with zero LLM API calls." Direct code inspection proves this is false \u2014 all 12 techniques import from llm_gateway.py and make real LLM API calls with template fallbacks.')

add_para(doc, 'The remaining work is organized into 6 weekly phases. Each phase represents one week of development effort and concludes with a mandatory testing gate that includes unit tests, integration tests, security regression tests, and performance baselines. Code may only be pushed after all tests in the weekly gate pass. This ensures no regressions are introduced and the system remains stable throughout the development process.')

add_para(doc, 'Current State: ~75% built. Backend AI engine is genuine with 12 real LLM-powered techniques. Auth security is hardened with JWT, httpOnly cookies, MFA, and role-based access. 7 dashboard pages (Tickets, Billing, Knowledge Base, Settings, AI Agents, Monitoring, Variants) are empty 15-line stubs. 6 critical bugs remain in the database and routing layers. 22 HIGH and 39 MEDIUM security findings are unresolved.')

add_para(doc, 'Target State: 100% production-ready. Zero CRITICAL/HIGH vulnerabilities. All DMD specifications implemented. FAKE Voting with multi-evaluator consensus scoring. CLARA RAG with HyDE, multi-query retrieval, and contextual compression. Full MAKER Framework pipeline with 6-24 LLM calls per query. Agent Lightning self-learning system. PostgreSQL Row-Level Security on all 93 tenant-scoped tables. 25 Loophole Solutions safety framework. All 7 dashboard pages functional. Full test suite with 80%+ coverage. Production deployment with monitoring, backups, and documentation.')

# ══════════════════════════════════════════
# 2. DAYS 1-3 CROSS-CHECK VERIFICATION
# ══════════════════════════════════════════
doc.add_heading('2. Days 1-3 Cross-Check Verification', level=1)

add_para(doc, 'Days 1-3 of the original roadmap were independently verified by reading every referenced file in the actual codebase. The following tables summarize the results. These items have been completed by the developer and are not part of the remaining work, except for the 6 items carried forward to Week 1.')

doc.add_heading('2.1 Day 1: Auth Security (13/13 DONE)', level=2)
add_para(doc, 'All 13 authentication and authorization security items from Day 1 are verified as DONE. The frontend uses jose library for JWT tokens with httpOnly cookies (not localStorage). Email verification is required before account activation. OTP comparison uses crypto.timingSafeEqual to prevent timing attacks. Password complexity enforces 8+ characters with upper, lower, and digit requirements. Dashboard API endpoints enforce JWT authentication returning 401 on failure. MFA uses a temporary session token with 5-minute TTL. Platform admin flag requires require_platform_admin. Billing status endpoint is authenticated. RAG is scoped to the JWT tenant only. MCP server enforces auth tokens. Chat widget validates company_id against the database. Chat API is authenticated with JWT verification returning 401 on failure.')

add_table(doc,
    ['ID', 'Item', 'Status'],
    [
        ['C-02', 'Frontend JWT tokens via jose library', 'DONE'],
        ['C-03', 'httpOnly cookies (not localStorage)', 'DONE'],
        ['H-03', 'Email verification required (is_verified: false)', 'DONE'],
        ['H-02', 'Timing-safe OTP (crypto.timingSafeEqual)', 'DONE'],
        ['M-20', 'Password complexity (8+ chars, upper, lower, digit)', 'DONE'],
        ['C-01', 'Dashboard API auth middleware (JWT + 401)', 'DONE'],
        ['C-09', 'MFA temp session token (5min TTL)', 'DONE'],
        ['C-10', 'Platform admin flag (require_platform_admin)', 'DONE'],
        ['C-11', 'Billing status endpoint authenticated', 'DONE'],
        ['C-12', 'RAG scoped to JWT tenant only', 'DONE'],
        ['C-04', 'MCP server auth token enforced', 'DONE'],
        ['H-14', 'Chat widget company_id validated against DB', 'DONE'],
        ['H-18', 'Chat API authenticated (JWT verify + 401)', 'DONE'],
    ],
    col_widths=[0.6, 4.5, 0.8]
)

doc.add_heading('2.2 Day 2: Data Protection (8/12 DONE)', level=2)
add_para(doc, 'Day 2 focused on data protection and configuration security. 8 of 12 items are fully DONE. The remaining 4 items (C-07, C-15, C-14, H-09) are partially done or not done and are carried forward to Week 1 Phase 1 as quick fixes. CORS wildcard has been removed. Config validators raise ValueError in production. X-Company-ID header is not trusted for tenant identification. PostgreSQL uses SSL with sslmode=require. Billing and admin routes are excluded from PUBLIC_PREFIXES. Workflow path parameters are validated. Redis password is required. SHA-256 is used for file integrity hashing.')

add_table(doc,
    ['ID', 'Item', 'Status', 'Gap'],
    [
        ['C-05', 'CORS wildcard removed', 'DONE', '\u2014'],
        ['C-06', 'Config validators raise ValueError in prod', 'DONE', '\u2014'],
        ['C-07', '.env.prod removed from git', 'NOT DONE', 'Still in git ls-files'],
        ['C-15', 'No default refresh pepper', 'PARTIAL', 'Prod RuntimeError; dev defaults empty'],
        ['C-08', 'X-Company-ID not trusted', 'DONE', '\u2014'],
        ['C-13', 'PostgreSQL SSL (sslmode=require)', 'DONE', '\u2014'],
        ['C-14', 'OAuth tokens encrypted at rest', 'PARTIAL', 'Fernet not AES-256-GCM'],
        ['H-05', 'Billing/admin not in PUBLIC_PREFIXES', 'DONE', '\u2014'],
        ['H-22', 'Workflow path param validated', 'DONE', '\u2014'],
        ['H-09', 'Pricing key from env var', 'PARTIAL', 'Hardcoded dev fallback remains'],
        ['H-10', 'Redis password required', 'DONE', '\u2014'],
        ['H-11', 'SHA-256 for file integrity', 'DONE', '\u2014'],
    ],
    col_widths=[0.5, 3.0, 0.8, 1.6]
)

doc.add_heading('2.3 Day 3: AI Rebuild (9/11 Met)', level=2)
add_para(doc, 'Day 3 focused on rebuilding the AI technique framework. 9 of 11 items are met. The two remaining items (AI-F01 and AI-F02) are carried forward to Week 1. A critical finding from the cross-check: the security audit claimed all 12 AI techniques were "pure regex with zero LLM API calls." Direct code inspection proves this is incorrect. All 12 techniques import from llm_gateway.py and make real LLM API calls, with template/fallback mechanisms for graceful degradation when LLM is unavailable. The table below shows the verification evidence for each technique, including the number of LLM call sites found in the actual source code.')

add_table(doc,
    ['ID', 'Item', 'Status', 'Evidence'],
    [
        ['AI-F01', 'Shared LLM client at techniques/llm_client.py', 'PARTIAL', 'llm_gateway.py exists at core/'],
        ['AI-F02', 'execute_with_llm() in base.py', 'NOT DONE', 'Method does not exist'],
        ['AI-01', 'Chain of Thought', 'REAL LLM', '2 LLM call sites + template fallback'],
        ['AI-02', 'Tree of Thought', 'REAL LLM', '1 LLM call site + heuristic fallback'],
        ['AI-03', 'ReAct', 'REAL LLM', '2 LLM call sites + pattern fallback'],
        ['AI-04', 'Reflexion', 'REAL LLM', '1 LLM call site + regex fallback'],
        ['AI-05', 'Self-Consistency', 'REAL LLM', '1 LLM call site + deterministic pipeline'],
        ['AI-06', 'Universe of Thought', 'REAL LLM', '1 LLM call site + weighted scoring'],
        ['AI-08', 'Least-to-Most', 'REAL LLM', '1 LLM call site + template decomposition'],
        ['AI-10', 'Step-Back', 'REAL LLM', '1 LLM call site + template broadening'],
        ['AI-11', 'Reverse Thinking', 'REAL LLM', '2 LLM call sites + template inversion'],
    ],
    col_widths=[0.5, 3.2, 0.8, 1.6]
)

doc.add_heading('2.4 Items Carried Forward to Week 1', level=2)
add_para(doc, 'The following 6 items from Days 2-3 were not fully resolved and are assigned to the beginning of Week 1 Phase 1 as quick fixes before proceeding with larger work items. These are estimated to take only 2-3 hours total and will be completed first in the phase.')

items_cf = [
    ('[C-07] ', 'Remove .env.prod from git tracking. The file is listed in .gitignore but was committed before the ignore entry was added. Run git rm --cached .env.prod, commit, and add a pre-commit hook preventing env files from being tracked.'),
    ('[C-15] ', 'Add a logging.warning() call that fires on startup when REFRESH_TOKEN_PEPPER is empty, regardless of environment. The production RuntimeError guard already exists; this adds development-time visibility.'),
    ('[C-14] ', 'Document the encryption decision with a brief ADR. Fernet (AES-128-CBC + HMAC-SHA256) provides authenticated encryption from the cryptography library. While not AES-256-GCM as originally specified, Fernet is a safe, audited implementation.'),
    ('[H-09] ', 'Add a production guard to PRICING_SIGNING_KEY in config.py that raises ValueError when ENVIRONMENT=production and the key still contains the dev default string, matching the pattern used by SECRET_KEY and other critical config values.'),
    ('[AI-F01] ', 'Create backend/app/core/techniques/llm_client.py as a thin re-export wrapper around the existing llm_gateway.py (632 lines). This satisfies the roadmap path requirement without duplicating implementation.'),
    ('[AI-F02] ', 'Add execute_with_llm() to BaseTechniqueNode in base.py that centralizes the common LLM call pattern: call llm_gateway.generate(), check for empty response, log token usage, return structured output.'),
]
for prefix, text in items_cf:
    add_bullet(doc, text, bold_prefix=prefix)

# ══════════════════════════════════════════
# 3. PHASE 1 (WEEK 1)
# ══════════════════════════════════════════
doc.add_heading('3. Phase 1: Security Carry-Forward + Security HIGH (Week 1)', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours across 1 week')
add_para(doc, 'Dependencies: Days 1-3 completed and cross-checked')
add_para(doc, 'Phase 1 closes out the remaining Day 2-3 carry-forward items, implements the FAKE Voting sub-system, builds the first dashboard page (Tickets), addresses all 22 HIGH severity security findings, and completes the Socket.io frontend integration. At the end of this week, zero CRITICAL and zero HIGH vulnerabilities should remain in the codebase. The week concludes with a comprehensive testing gate covering unit tests, integration tests, and security regression tests before any code is pushed.')

doc.add_heading('3.1 Carry-Forward Quick Fixes (2-3 hours)', level=2)
add_para(doc, 'These 6 items are the unfinished work from Days 2-3. They are quick fixes estimated at 2-3 hours total and must be completed first before any larger work begins. Each fix is small and well-defined, making them ideal for immediate resolution. See Section 2.4 above for detailed descriptions of each item.')

doc.add_heading('3.2 FAKE Voting Sub-System (4-5 hours)', level=2)
add_para(doc, '[GAP-FAKE] The MAKER Framework validator (11_maker_validator.py, 684 lines) generates K solutions and scores them but lacks the FAKE Voting mechanism specified in the DMD. This is a core component of the MAKER (Maximal Agentic Decomposition + FAKE + Red-Flagging) framework that ensures AI responses are high quality and safe. The implementation must generate 3-5 candidate responses per customer query, score each using multiple evaluation dimensions (fluency, relevance, safety, brand voice compliance), apply Red-Flagging for hallucinations and policy violations, and produce a weighted consensus score to select the best response. The FAKE Voting results must be logged for the Agent Lightning self-learning system to use as training data.')

doc.add_heading('3.3 Dashboard: Tickets Page (4 hours)', level=2)
add_para(doc, '[GAP-FE1] Build the ticket management page, which is currently a 15-line placeholder component. This page must implement a comprehensive ticket list view with sortable columns (ID, customer name, subject, status, priority, assigned agent, created date), a filter panel supporting status, priority, agent, and date range filters, bulk action capabilities (assign, close, escalate), a detail drawer showing full ticket conversation history, and real-time status updates via the Socket.io integration built later in this phase. The page must connect to real backend API endpoints and display actual ticket data from the database.')

doc.add_heading('3.4 Security HIGH - Middleware and Webhooks (8-10 hours)', level=2)
add_para(doc, 'This section addresses the most critical security findings from the audit. Each item is classified as HIGH severity and must be resolved before the codebase can be considered production-ready. The items are grouped into three categories: middleware and access control, webhook and CSRF security, and additional HIGH fixes.')

doc.add_heading('3.4.1 Middleware and Access Control', level=3)
mid_items = [
    ('[H-01] Open Redirect Fix: ', 'Validate redirect query parameter against whitelist of allowed paths. Block double-encoding tricks and protocols other than single forward slash paths. File: src/lib/auth-cookies.ts.'),
    ('[H-04] Content-Security-Policy: ', 'Add comprehensive CSP to security headers middleware. Policy: default-src \'self\', script-src \'self\' \'nonce-{random}\', style-src \'self\' \'unsafe-inline\' (required for Tailwind), img-src \'self\' data: https:, connect-src \'self\' wss: https:. File: backend/app/middleware/security_headers.py.'),
    ('[H-06] IP Extraction Standardization: ', 'Create shared get_client_ip() utility in backend/app/core/utils.py respecting TRUSTED_PROXY_COUNT configuration. Apply consistently to ip_allowlist.py, request_logger.py, and rate_limit.py to eliminate inconsistent IP extraction logic.'),
    ('[H-13] Billing Role Restrictions: ', 'Add owner/admin role checks to cancel subscription, process refund, and modify billing endpoints. Only users with owner or admin role should be able to perform billing modifications. File: backend/app/api/billing.py.'),
    ('[H-21] Auth Rate Limiting: ', 'Add Redis sliding window rate limits: Login 5 attempts/min/email, Register 3 attempts/min/IP, OTP 10 attempts/min/phone, Password Reset 3 attempts/hour/email. Use Redis pipeline for atomic check-and-increment. File: backend/app/middleware/rate_limit.py.'),
]
for prefix, text in mid_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('3.4.2 Webhook and CSRF Security', level=3)
wh_items = [
    ('[H-07] Webhook Signature Verification: ', 'Change webhook processing logic to reject requests when webhook_secret is missing (fail-closed security posture), rather than silently accepting all payloads. File: backend/app/api/billing_webhooks.py.'),
    ('[H-08] Webhook Replay Protection: ', 'Add timestamp validation (reject requests older than 5 minutes) and store processed webhook IDs in Redis with TTL for all webhook handlers (Paddle, Twilio, Brevo, Shopify). Files: backend/app/api/billing_webhooks.py, sms_channel.py, notification_webhooks.py.'),
    ('[H-19] CSRF Protection: ', 'Generate server-side CSRF tokens stored in httpOnly cookies. Validate on all POST/PUT/DELETE routes using cookie-based auth. For Bearer token auth, enforce CSRF on cookie-auth endpoints only. File: backend/app/middleware/csrf.py.'),
    ('[H-20] Mock Login Production Gate: ', 'Disable mock login endpoint entirely when ENVIRONMENT=production. Add startup check and runtime guard. File: dashboard/src/app/api/auth/login/route.ts.'),
]
for prefix, text in wh_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('3.4.3 Additional HIGH Fixes', level=3)
add_items = [
    ('[H-12] Google OAuth Token Exchange: ', 'Move token from URL query parameter to POST body to prevent token appearing in browser history, server logs, and referrer headers. File: backend/app/services/auth_service.py.'),
    ('[H-16] HTML Injection in Email Templates: ', 'Sanitize customer names and AI responses with bleach/markupsafe before HTML interpolation in email templates. File: dashboard/src/lib/notifications.ts.'),
    ('[H-17] Channel Status API Key Leakage: ', 'Remove first 8 characters of Brevo API key and first 6 characters of Twilio SID from channel status API response. Return only "configured" or "not configured" status. File: dashboard/src/app/api/channels/status/route.ts.'),
]
for prefix, text in add_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('3.5 Remaining HIGH Fixes + Socket.io Frontend (6-8 hours)', level=2)

doc.add_heading('3.5.1 Remaining HIGH Items', level=3)
rem_items = [
    ('[H-15] Webhook Management Auth: ', 'Add authentication to webhook event query and retry endpoints. Only admin users can access these management endpoints. File: backend/app/api/webhooks.py.'),
    ('[M-27] User Enumeration Prevention: ', 'Make check-email endpoint return identical response for existing and non-existing emails to prevent account enumeration. Add rate limiting. File: src/app/api/auth/check-email/route.ts.'),
    ('[M-37] Hardcoded SMS Phone: ', 'Fix notification service sending SMS to +1234567890 instead of the actual customer phone number. Route SMS sending through backend API. File: dashboard/src/lib/notifications.ts.'),
    ('[D-02] Analytics Mock Data: ', 'Connect dashboard analytics components to real backend APIs. Currently the analytics module silently returns fake numbers. File: dashboard/src/lib/analytics-api.ts.'),
    ('[D-03] Email Content Sanitization: ', 'Route email sending through backend API instead of calling Brevo directly from the frontend, which exposes API keys. File: dashboard/src/app/api/send-email/route.ts.'),
]
for prefix, text in rem_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('3.5.2 Socket.io Frontend Integration', level=3)
add_para(doc, '[GAP-SOCKET] The backend has a full Socket.io server implementation (backend/app/core/socketio.py) but the frontend has zero Socket.io client code. This is a significant gap because PARWA\'s value proposition includes real-time AI-powered customer care. The implementation must: (1) Install socket.io-client npm package. (2) Create a shared useSocket React hook managing connection lifecycle, reconnection, and event subscriptions. (3) Implement real-time ticket status updates across all dashboard pages. (4) Add live chat message streaming for the customer-facing widget. (5) Push notification events (new ticket, escalation, SLA breach warnings) to the dashboard. (6) Implement proper authentication token passing during socket handshake. (7) Handle reconnection with exponential backoff and event replay for missed messages during disconnection.')

doc.add_heading('3.6 Week 1 Testing Gate', level=2)
add_para(doc, 'Before any code from Week 1 is pushed to the repository, ALL of the following tests must pass. This is a mandatory gate that cannot be bypassed. If any test fails, the corresponding code must be fixed before pushing.')

test_w1 = [
    'Unit tests for all 6 carry-forward fixes (6 test files covering C-07, C-15, C-14, H-09, AI-F01, AI-F02)',
    'Unit tests for FAKE Voting mechanism: candidate generation, multi-evaluator scoring, Red-Flagging, weighted consensus',
    'Unit tests for Tickets page components: list rendering, sort, filter, bulk actions, detail drawer',
    'Integration tests for all HIGH security fixes: middleware headers, webhook verification, CSRF tokens, rate limiting',
    'Integration test for Socket.io connection lifecycle: connect, authenticate, receive events, reconnect',
    'Security regression test suite: verify zero new CRITICAL or HIGH vulnerabilities introduced',
    'Coverage target: 80%+ on all modified files',
]
for t in test_w1:
    add_bullet(doc, t)

p = doc.add_paragraph()
r = p.add_run('GATE: All tests pass \u2192 git push \u2192 proceed to Week 2')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x0A, 0x82, 0x8A)

doc.add_heading('3.7 Phase 1 Acceptance Criteria', level=2)
add_para(doc, 'All 6 Day 2-3 carry-forward items resolved and verified. FAKE Voting generates K candidates with multi-evaluator consensus scoring and Red-Flagging for hallucinations. Tickets page displays real data from backend APIs with filters, sort, and bulk actions. Socket.io provides real-time updates for tickets, chat, and notifications. Zero CRITICAL vulnerabilities remain. All 22 HIGH severity findings resolved. All tests in Week 1 gate pass with 80%+ coverage.')

# ══════════════════════════════════════════
# 4. PHASE 2 (WEEK 2)
# ══════════════════════════════════════════
doc.add_heading('4. Phase 2: AI Frameworks + Dashboard Batch 1 (Week 2)', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours across 1 week')
add_para(doc, 'Dependencies: Phase 1 complete (security baseline + FAKE Voting + Socket.io)')
add_para(doc, 'Phase 2 rebuilds the higher-level AI systems that compose individual techniques into a complete customer care pipeline. This includes CLARA RAG with advanced retrieval capabilities, the full MAKER Framework pipeline, real DSPy integration, Agent Lightning self-learning, 3-tier hybrid optimization, and critical Smart Router bug fixes. After this phase, the AI engine will perform genuine multi-stage reasoning for every customer query, not just single-technique responses.')

doc.add_heading('4.1 CLARA RAG Rebuild (8-10 hours)', level=2)
add_para(doc, '[AI-14] The CLARA RAG system is currently a scoring wrapper without the advanced retrieval capabilities specified in the DMD document. This rebuild adds three key features that transform CLARA from a basic retriever into an enterprise-grade knowledge system capable of handling complex customer care queries across multiple knowledge domains.')

clara_items = [
    ('HyDE (Hypothetical Document Embedding): ', 'When a customer query arrives, use the LLM to generate a hypothetical answer document. Embed this hypothetical answer and use it as the primary retrieval vector instead of the raw query. This dramatically improves retrieval quality because the embedding space is optimized for document embeddings, not query embeddings. The hypothetical answer captures semantic intent that keyword matching misses entirely.'),
    ('Multi-Query Retrieval: ', 'Generate 3 alternative phrasings of the original query using the LLM, then retrieve documents for all 4 queries (original + 3 alternatives). Merge and deduplicate the results, then rank by aggregate relevance score. This handles vocabulary mismatch between customer language and knowledge base terminology, which is a common failure mode in customer care systems.'),
    ('Contextual Compression: ', 'After initial retrieval, use the LLM to extract only the relevant sentences from each document chunk. This reduces context window usage by 50-70% and improves accuracy by eliminating noise from irrelevant parts of retrieved documents. Critical for staying within LLM token limits while maximizing useful context. Implementation: backend/app/core/rag/clara_rag.py.'),
]
for prefix, text in clara_items:
    add_bullet(doc, text, bold_prefix=prefix)

add_para(doc, '[AI-14b] RAG Re-ranking with LLM: After initial retrieval and deduplication, pass each document plus the original query to the LLM for relevance scoring on a 1-10 scale. Sort all retrieved documents by score and use only the top-K for context generation. This two-stage retrieval (vector similarity + LLM re-ranking) significantly improves precision. File: backend/app/core/rag/reranker.py.')

doc.add_heading('4.2 MAKER Framework Full Pipeline (8-10 hours)', level=2)
add_para(doc, '[AI-15] The MAKER Framework should execute 6-24 LLM calls per customer query, depending on the tier (Mini PARWA, PARWA, PARWA High). Currently many LangGraph nodes are thin wrappers that pass through without real processing. This rebuild implements the full multi-stage pipeline as specified in the DMD:')

maker_items = [
    ('Map stage (Node 01): ', 'LLM classifies query type (refund, technical, billing, complaint, FAQ, general) to determine specialized agent routing. This classification drives the entire pipeline flow.'),
    ('Analyze stage (Nodes 02-03): ', 'Empathy detection via LLM analyzing customer emotional state. Sentiment scoring with structured output. Urgency level detection. Entity extraction for order numbers, product names, dates, and amounts.'),
    ('Knowledge stage (Node 04): ', 'CLARA RAG retrieval using HyDE + Multi-Query + Compression (built in 4.1) for knowledge base document retrieval and context assembly.'),
    ('Generate stage (Nodes 06-10): ', 'Route to specialized agents (Refund, Technical, Billing, Complaint, FAQ) using the appropriate AI technique per agent type. Each agent uses its best-suited technique from the 12 available.'),
    ('Evaluate stage (Node 11): ', 'FAKE Voting (built in Phase 1) generates K candidate responses, applies multi-evaluator scoring across 4+ dimensions, runs Red-Flagging for hallucinations and policy violations, produces weighted consensus score.'),
    ('Control stage (Node 12): ', 'Jarvis Control System monitors response quality, checks for policy compliance, verifies tone and brand voice. Overrides or escalates when quality falls below configurable threshold.'),
    ('Guard stage (Node 14): ', 'Block harmful content, enforce brand voice guidelines, check for off-topic responses, verify response length is appropriate for channel, ensure regulatory compliance (GDPR, PCI-DSS).'),
    ('Delivery stage (Node 15): ', 'Channel-specific formatting: SMS (160 chars), HTML email, TTS voice markup, quick-reply chat buttons. Each channel has different constraints and formatting requirements.'),
]
for prefix, text in maker_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('4.3 DSPy Integration + Agent Lightning (6-8 hours)', level=2)
add_para(doc, '[AI-12] Replace the existing StubModule with real DSPy integration. The DSPy framework provides a programmatic interface for composing LLM calls into optimized pipelines. Implementation must include: (1) DSPy Signature definition for customer care (input: query + conversation_context + knowledge_context, output: response + confidence + escalation_flag). (2) DSPy Module chaining CoT and ReAct techniques into a unified pipeline. (3) DSPy teleprompt optimization using the Bayesian Signature Optimizer to automatically tune prompt templates based on historical conversation data. (4) Integration with the Smart Router for tier-aware technique selection. File: backend/app/core/dspy/.')

add_para(doc, '[AI-13 / GAP-AGENT-LIGHTNING] Rebuild the Agent Lightning self-learning system. Database models already exist (TrainingDataset, TrainingCheckpoint, AgentMistake, AgentPerformance, TrainingRun) but no training pipeline implementation exists. Build: (1) Mistake collection pipeline that logs every AI response flagged by FAKE Voting Red-Flagging or Jarvis Control System override. (2) Training data preparation that converts mistakes into LLaMA-3-8B fine-tuning format via Unsloth. (3) Google Colab training notebook that executes fine-tuning with collected data. (4) Model deployment pipeline that replaces the default technique model with the fine-tuned version. (5) A/B testing framework to compare fine-tuned vs default model performance before full rollout.')

doc.add_heading('4.4 3-Tier Hybrid + Smart Router Fixes (6-8 hours)', level=2)
add_para(doc, '[AI-16] Implement 3-tier optimization routing as specified in the DMD. Each pricing tier gets different AI processing depth: Tier 1 (Mini PARWA, $1K/mo) uses the fastest single AI technique for simple queries. Tier 2 (PARWA, $2.5K/mo) uses the best single technique based on intent classification. Tier 3 (PARWA High, $4K/mo) uses the full MAKER Framework pipeline with all 8 stages and FAKE Voting. The Smart Router must determine the tier from the company subscription and apply the appropriate processing depth automatically.')

add_para(doc, '[BUG-HEALTHTRACKER] Smart Router HealthTracker currently stores provider health state in Python dictionaries. With multiple Celery workers, each worker has independent state, leading to inconsistent routing decisions. Refactor to Redis-backed state with atomic operations using Redis pipelines. This ensures all workers see the same provider health status and make consistent routing decisions across the entire system.')

add_para(doc, '[BUG-MOCKVECTOR] The RAG system currently defaults to MockVectorStore which uses SHA-256 pseudo-embeddings, producing terrible retrieval quality. Reverse the default logic: use PgVectorStore when PostgreSQL + pgvector is available (always in production, usually in development). Fall back to MockVectorStore only when explicitly configured or when pgvector is truly unavailable. This single change dramatically improves RAG quality without any algorithm changes.')

doc.add_heading('4.5 Week 2 Testing Gate', level=2)
test_w2 = [
    'Unit tests for CLARA RAG: HyDE generation, multi-query retrieval, contextual compression, LLM re-ranking',
    'Unit tests for MAKER pipeline: all 8 stages tested independently with mock LLM responses',
    'Unit tests for DSPy integration: Signature validation, Module chaining, optimizer compilation',
    'Unit tests for Agent Lightning: mistake collection, data preparation, A/B test framework',
    'Integration test: full customer query through all MAKER pipeline stages with real LLM calls',
    'Integration test: Smart Router tier selection per variant (Mini/PARWA/PARWA High)',
    'Integration test: Agent Lightning training collection and model deployment workflow',
    'Load test: 10 concurrent conversations through full pipeline, document performance baselines',
    'Coverage target: 80%+ on all modified files',
]
for t in test_w2:
    add_bullet(doc, t)

p = doc.add_paragraph()
r = p.add_run('GATE: All tests pass \u2192 git push \u2192 proceed to Week 3')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x0A, 0x82, 0x8A)

doc.add_heading('4.6 Phase 2 Acceptance Criteria', level=2)
add_para(doc, 'CLARA RAG performs real HyDE generation, multi-query retrieval, contextual compression, and LLM re-ranking. MAKER Framework executes 6+ LLM calls per query through all 8 stages of the pipeline. DSPy uses real Signatures (no StubModule) with optimized teleprompt. Agent Lightning collects mistakes from FAKE Voting and Jarvis overrides, triggers training, and deploys models. 3-tier routing applies correct processing depth per subscription variant. Smart Router uses Redis-backed health state. RAG uses PgVectorStore by default. All Week 2 gate tests pass.')

# ══════════════════════════════════════════
# 5. PHASE 3 (WEEK 3)
# ══════════════════════════════════════════
doc.add_heading('5. Phase 3: Security MEDIUM + Dashboard + LangGraph (Week 3)', level=1)
add_para(doc, 'Estimated Effort: 40-45 hours across 1 week')
add_para(doc, 'Dependencies: Phase 2 complete (AI frameworks working)')
add_para(doc, 'Phase 3 is the heaviest week in terms of item count. It addresses all 39 MEDIUM security findings, completes all 7 dashboard pages (6 remaining after Tickets from Week 1), and finishes all 19 LangGraph pipeline nodes. This phase transforms PARWA from a working backend with empty frontend stubs into a fully functional application with polished UI and comprehensive security.')

doc.add_heading('5.1 Security MEDIUM Fixes (8-10 hours)', level=2)
add_para(doc, 'The following 17 MEDIUM severity findings must be addressed. While individually less critical than HIGH findings, collectively they represent significant attack surface. Each fix is described with its specific implementation requirement and target file.')

medium_items = [
    ('[M-01] ', 'Remove user_role from AuthorizationError details. Return generic "access denied" message to prevent role enumeration. File: backend/app/api/deps.py.'),
    ('[M-04] ', 'Replace simple "@" email validation with email-validator package supporting RFC 5322 compliance. File: backend/app/api/verification.py.'),
    ('[M-05] ', 'Rate limiter fail-closed: block all requests when Redis is down, not allow all. Log Redis failure for monitoring. File: backend/app/middleware/rate_limit.py.'),
    ('[M-06] ', 'API key authentication must return 401 when no header present, not pass through to next handler. File: backend/app/middleware/api_key_auth.py.'),
    ('[M-08] ', 'Add explicit authentication dependency to events API endpoint. File: backend/app/main.py.'),
    ('[M-11] ', 'Add Cache-Control: no-store, no-cache, must-revalidate headers to all authentication responses. File: backend/app/middleware/security_headers.py.'),
    ('[M-13] ', 'Replace setattr-based user update in admin.py with explicit field whitelist assignment. File: backend/app/api/admin.py.'),
    ('[M-14] ', 'Add Pydantic request models to ai_engine endpoints that currently accept raw dict bodies. File: backend/app/api/ai_engine.py.'),
    ('[M-16] ', 'Add Twilio request signature verification to SMS status callback endpoint. File: backend/app/api/sms_channel.py.'),
    ('[M-17] ', 'Replace str(e) exception messages with generic errors in knowledge base API responses. File: backend/app/api/knowledge_base.py.'),
    ('[M-19] ', 'Fix visitor token verification that passes on exception instead of rejecting the request. File: backend/app/api/chat_widget.py.'),
    ('[M-23] ', 'Fix MCP server CORS configuration that falls back to ["*"] on exception. File: mcp_server/main.py.'),
    ('[M-26] ', 'Add security headers (X-Content-Type-Options, X-Frame-Options, HSTS, CSP) to Next.js API routes via middleware. File: src/middleware.ts.'),
    ('[M-28] ', 'Sanitize email content before passing to Brevo API to prevent injection. File: dashboard/src/app/api/send-email/route.ts.'),
    ('[M-32] ', 'Add max payload size validation (1MB limit) to all Celery task definitions. File: backend/app/tasks/celery_app.py.'),
    ('[M-33] ', 'Escape % and _ characters in search queries before ILIKE to prevent SQL wildcard injection. Files: backend/app/api/admin.py, tickets.py.'),
    ('[M-35] ', 'Add role check to notification send endpoint. Only admin/manager roles can send notifications. File: backend/app/api/notifications.py.'),
]
for prefix, text in medium_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('5.2 Dashboard Pages Batch 2 (12-14 hours)', level=2)
add_para(doc, 'Six dashboard pages need to be built from their current 15-line placeholder state into fully functional pages connected to real backend APIs. Each page must support real-time updates via Socket.io, proper loading states, error handling, and responsive design.')

dash_items = [
    ('[GAP-FE2] Billing Dashboard Page: ', 'Subscription overview with current plan details, renewal date, and usage meters. Invoice list with PDF download capability. Payment method management (add, remove, set default). Upgrade/downgrade flow via Paddle integration with prorated amount display. Usage analytics per billing period.'),
    ('[GAP-FE3] Knowledge Base Dashboard Page: ', 'Document upload with drag-and-drop support and progress indicators. Document list with status indicators (indexed, indexing, failed). Full-text search across all articles. Manual reindex triggering. Document deletion with confirmation. Bulk upload support.'),
    ('[GAP-FE4] Settings Dashboard Page: ', 'Tabbed interface with 5 sections: Company Profile (name, logo, industry, timezone), API Keys (generate, view, revoke with copy-to-clipboard), Team Management (invite members, assign roles, remove), Notification Preferences (email, SMS, push toggles per category), Security (MFA setup, active sessions, IP allowlist configuration).'),
    ('[GAP-FE5] AI Agents Dashboard Page: ', 'Agent list with real-time status (active, training, idle). Performance metrics per agent (resolution rate, average response time, customer satisfaction score). Capability matrix showing which AI techniques each agent uses. Training history from Agent Lightning. Configuration panel for agent-specific settings.'),
    ('[GAP-FE6] Monitoring Dashboard Page: ', 'System health panel (API latency P50/P95/P99, error rates, queue depths). AI engine metrics (token usage per request, confidence score distribution, technique usage counts). Real-time alert feed with severity levels. SLA compliance tracking with breach warnings.'),
    ('[GAP-FE7] Variants Dashboard Page: ', 'Current variant tier display with feature comparison matrix (170+ features from Feature Registry). Capability toggles to enable/disable features per tier. Workload distribution visualization showing queries per tier. Token budget configuration per tier and per customer.'),
]
for prefix, text in dash_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('5.3 LangGraph Node Completion (12-14 hours)', level=2)
add_para(doc, 'All 19 LangGraph nodes must perform real processing with no pass-through stubs. The following 12 nodes need completion or rebuilding. Each node must integrate with the AI systems built in Phase 2 (CLARA RAG, MAKER pipeline, DSPy, Smart Router) and use real LLM calls via llm_gateway.py.')

node_items = [
    ('[LG-01] Node 02 - Empathy Engine: ', 'LLM analyzes customer emotional state (anger, frustration, confusion, satisfaction, urgency) and adjusts response tone accordingly. Uses structured output for emotion classification.'),
    ('[LG-02] Node 04 - Base Domain Agent: ', 'Connect to CLARA RAG for knowledge retrieval and use Smart Router technique selection for response generation.'),
    ('[LG-03] Node 05 - FAQ Agent: ', 'Embedding similarity matching against FAQ knowledge base (not just keyword search). LLM fallback when no match found above confidence threshold.'),
    ('[LG-04] Node 06 - Refund Agent: ', 'Connect to Paddle API for refund eligibility verification, amount calculation, and initiation through the billing system.'),
    ('[LG-05] Node 07 - Technical Agent: ', 'Use ReAct technique with diagnostic tools. Check service health, search knowledge base for known issues, guide troubleshooting step-by-step.'),
    ('[LG-06] Node 08 - Billing Agent: ', 'Connect to billing system for subscription queries, invoice lookups, payment troubleshooting, plan changes via Paddle API.'),
    ('[LG-07] Node 09 - Complaint Agent: ', 'Emotional intelligence with sentiment analysis, service recovery playbooks, automatic escalation when severity exceeds AI confidence threshold.'),
    ('[LG-08] Node 10 - Escalation Agent: ', 'Determine escalation necessity based on topic, sentiment, confidence, and conversation length. Route to correct human queue with full context transfer.'),
    ('[LG-09] Node 12 - Control System: ', 'Jarvis monitoring for response accuracy, policy compliance, brand voice, PII leakage. Override or escalate when quality falls below configurable threshold.'),
    ('[LG-10] Node 13 - DSPy Optimizer: ', 'Connect to rebuilt DSPy module for technique selection optimization based on historical performance data.'),
    ('[LG-11] Node 14 - Guardrails: ', 'Block harmful content, enforce brand voice, check for off-topic responses, verify length, ensure regulatory compliance.'),
    ('[LG-12] Node 15 - Channel Delivery: ', 'SMS truncation (160 chars), HTML email formatting, TTS voice markup, quick-reply chat button generation per channel type.'),
]
for prefix, text in node_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('5.4 Week 3 Testing Gate', level=2)
test_w3 = [
    'Unit tests for all 17 MEDIUM security fixes (individual test per fix)',
    'Unit tests for all 6 dashboard pages: component rendering, API integration, filter/sort, error states',
    'Unit tests for all 12 LangGraph nodes: input validation, LLM call structure, response parsing, error handling',
    'Integration test: full 19-node LangGraph pipeline execution with real LLM calls at each stage',
    'Integration test: all 7 dashboard pages loading real data from backend APIs simultaneously',
    'Multi-tenant isolation test: verify no cross-tenant data leakage across all API endpoints',
    'Coverage target: 80%+ on all modified files',
]
for t in test_w3:
    add_bullet(doc, t)

p = doc.add_paragraph()
r = p.add_run('GATE: All tests pass \u2192 git push \u2192 proceed to Week 4')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x0A, 0x82, 0x8A)

# ══════════════════════════════════════════
# 6. PHASE 4 (WEEK 4)
# ══════════════════════════════════════════
doc.add_heading('6. Phase 4: RLS + Critical Bugs + Infrastructure (Week 4)', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours across 1 week')
add_para(doc, 'Dependencies: Phase 3 complete (all nodes working, dashboards built)')
add_para(doc, 'Phase 4 implements defense-in-depth database security through PostgreSQL Row-Level Security, fixes 4 critical deployment bugs that would cause failures in production, and hardens infrastructure for production deployment. This is the security hardening phase that makes PARWA enterprise-ready.')

doc.add_heading('6.1 Critical Database Bug Fixes (10-12 hours)', level=2)
add_para(doc, 'These 4 bugs are classified as critical because they would cause data corruption, incorrect behavior, or data loss in a production deployment. Each must be fixed and tested thoroughly.')

bug_items = [
    ('[BUG-1] Migration Chain FK Mismatch: ', 'Alembic migrations 003 and 005 reference sessions.id as a foreign key, but the table was renamed to tickets (BL01) without updating the migration FK references. This causes alembic upgrade to fail on fresh databases. Create migration 020 that handles both cases: databases that ran the old migrations (rename FK references) and fresh databases (use correct references from start). File: database/alembic/versions/020_fix_session_ticket_fk.py.'),
    ('[BUG-2] UUID Type Inconsistency: ', 'OutboundEmail and EmailDeliveryEvent models use UUID(as_uuid=True) for primary keys while all other 95 models use String(36). This causes join failures and type mismatch errors when these tables are queried alongside other models. Convert the 2 outlier models to String(36) with the _uuid() helper function for consistency. Create a data migration preserving existing data.'),
    ('[BUG-3] Route Prefix Mismatch: ', 'variant_check.py checks for /api/v1/tickets but tickets.py registers routes with /tickets prefix (without /api/v1). Variant rate limiting and feature restrictions never actually enforce on ticket creation endpoints because the prefix never matches. Align the prefix in variant_check.py with the actual route registration in tickets.py.'),
    ('[BUG-4] Jarvis Session Persistence: ', 'Jarvis sessions are stored in an in-memory Python Map(), meaning all session state is lost every time the server restarts. This is unacceptable for production. Migrate to Redis-backed session storage with configurable TTL and automatic cleanup. The backend already has JarvisSession database models that can be leveraged for persistence.'),
]
for prefix, text in bug_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('6.2 PostgreSQL Row-Level Security (10-12 hours)', level=2)
add_para(doc, '[GAP-RLS] PostgreSQL Row-Level Security provides defense-in-depth alongside the existing application-level tenant isolation via SQLAlchemy middleware. Even with perfect application code, RLS protects against SQL injection, ORM bypass, direct database access, and future code bugs that might accidentally expose cross-tenant data. This is an enterprise security requirement for multi-tenant SaaS applications.')

rls_items = [
    'Create Alembic migration enabling RLS on all 93 tenant-scoped tables (tables containing company_id column)',
    'For each table: CREATE POLICY restricting SELECT, UPDATE, INSERT, DELETE operations to rows where company_id matches the current tenant context',
    'Implement PostgreSQL function that extracts company_id from the JWT token passed as session variable via SET app.current_tenant_id',
    'Modify the database connection factory to execute SET app.current_tenant_id on every connection checkout from the pool',
    'Ensure a RLS bypass role exists for Alembic migrations and admin operations that require cross-tenant data access',
    'Write comprehensive tests: direct database queries without app.current_tenant_id must return zero rows. Admin operations via superuser role must work correctly.',
    'Files: database/alembic/versions/ (new migration), database/rls_policies.sql (new policy definitions)',
]
for t in rls_items:
    add_bullet(doc, t)

doc.add_heading('6.3 Infrastructure Hardening (8-10 hours)', level=2)
add_para(doc, 'Infrastructure hardening covers 13 items from the LOW severity findings plus Docker security and environment management. While individually lower severity, they are essential for production deployment readiness.')

infra_items = [
    ('[L-01] JWT RS256 Migration: ', 'Migrate from HS256 (symmetric) to RS256 (asymmetric) allowing frontend token verification without shared signing key. Generate RSA key pair, configure backend signing and frontend verification.'),
    ('[L-02] Token Blacklist: ', 'Redis-backed token blacklist checking jti claims on every authenticated request. Celery beat cleanup task every 6 hours for expired entries.'),
    ('[L-04] Rate Limit Cleanup: ', 'Periodic Celery beat task (every 6 hours) removing expired rate limit entries from Redis to prevent memory bloat.'),
    ('[L-05] Circuit Breaker Thread Safety: ', 'Add threading.Lock to circuit breaker state machine for thread-safe state transitions in multi-threaded FastAPI workers.'),
    ('[L-09] Async Login Endpoint: ', 'Convert login endpoint from synchronous to async using async database session for better concurrency.'),
    ('[L-11] File Magic-Byte Validation: ', 'Add magic-byte validation for uploaded files (PNG: 0x89504E47, PDF: 0x25504446, JPEG: 0xFFD8FF). Do not trust Content-Type header alone.'),
    ('[L-12] JWT Key Rotation: ', 'Implement key rotation mechanism supporting multiple active signing keys with versioned metadata in Redis. Graceful rotation without invalidating existing sessions.'),
    ('[INF-01] Docker Security: ', 'All production containers run as non-root user. Remove unnecessary packages from images. Set CPU and memory limits. Enable read-only root filesystem. Configure health checks for all services.'),
    ('[INF-02] Environment Variable Audit: ', 'Create .env.production.template with all required variables, descriptions, and examples. Add startup validation to fail fast on missing required variables.'),
    ('[M-24] Dev Docker Port Binding: ', 'Restrict database (5432) and Redis (6379) ports to 127.0.0.1 binding. Only backend (8000) and frontend (3000) should be externally accessible.'),
    ('[M-25] Redis Dev Password: ', 'Add default Redis password even in development environment to prevent accidental exposure.'),
    ('[M-38] Google AI Key in URL: ', 'Move Google AI key from URL query parameter to request header in chat API route.'),
]
for prefix, text in infra_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('6.4 Week 4 Testing Gate', level=2)
test_w4 = [
    'Unit tests for all 4 critical bug fixes (migration FK, UUID consistency, route prefix, session persistence)',
    'Integration test: fresh database deployment with alembic upgrade head succeeds without errors',
    'Integration test: RLS enforcement - direct SQL queries return zero rows without tenant context',
    'Integration test: multi-tenant isolation via both RLS and application middleware (dual-layer verification)',
    'Security test: JWT RS256 verification, token blacklist check, key rotation without session invalidation',
    'Infrastructure test: Docker containers start as non-root, resource limits enforced, health checks pass',
    'Coverage target: 80%+ on all modified files',
]
for t in test_w4:
    add_bullet(doc, t)

p = doc.add_paragraph()
r = p.add_run('GATE: All tests pass \u2192 git push \u2192 proceed to Week 5')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x0A, 0x82, 0x8A)

# ══════════════════════════════════════════
# 7. PHASE 5 (WEEK 5)
# ══════════════════════════════════════════
doc.add_heading('7. Phase 5: Loopholes + Incomplete Features + Full Testing (Week 5)', level=1)
add_para(doc, 'Estimated Effort: 40-45 hours across 1 week')
add_para(doc, 'Dependencies: Phase 4 complete (RLS active, critical bugs fixed, infrastructure hardened)')
add_para(doc, 'Phase 5 builds the 25 Loophole Solutions safety framework (a core PARWA differentiator), completes all remaining incomplete features, and runs the most comprehensive testing in the entire roadmap. This is the quality assurance phase that proves PARWA is production-ready.')

doc.add_heading('7.1 25 Loophole Solutions Framework (10-12 hours)', level=2)
add_para(doc, '[GAP-LOOPHOLE] The DMD defines 25 specific production loophole scenarios where AI could cause harm, leak data, or make incorrect decisions. These represent real-world failure modes observed in AI customer care systems. Currently no centralized implementation exists. Build LoopholeRegistry (dataclass mapping each loophole to severity, detection pattern, and countermeasure), LoopholeDetectionEngine in backend/app/core/loophole_engine.py that scans AI responses for loophole patterns using both rule-based checks and LLM-based analysis, and integrate into LangGraph guardrails (Node 14) as a mandatory check before any response is delivered to a customer.')

add_para(doc, 'The 25 loophole categories include: hallucination of policies/discounts that do not exist, promising refunds outside policy, sharing one customer\'s data with another, confirming account details without verification, generating legal/medical/financial advice, reproducing copyrighted content, social engineering susceptibility, PII leakage in responses, biased responses based on customer demographics, emotional manipulation, off-topic responses, contradictory statements within same conversation, numerical errors in calculations, timezone/date errors, currency conversion mistakes, escalation failures for high-risk queries, brand voice violations, competitive information disclosure, SLA commitment over-promises, regulatory non-compliance (GDPR, PCI-DSS, HIPAA), unauthorized data access attempts, conversation state confusion across sessions, infinite loop triggers, and resource exhaustion attacks through complex queries.')

doc.add_heading('7.2 Incomplete Feature Completion (10-12 hours)', level=2)
add_para(doc, 'Six features identified as incomplete in the codebase audit need to be fully implemented. These are features where the infrastructure exists but the core logic is missing or uses placeholder implementations.')

inc_items = [
    ('[IC-01] Voice Server LLM Intent: ', 'Replace keyword-based intent detection in parwa_voice_server.py with real LLM-based classification via Smart Router. Maintain existing Twilio call handling infrastructure, only replace the intent classification component.'),
    ('[IC-02] Sentiment NLP: ', 'Replace keyword-based sentiment analysis with LLM-powered structured output producing sentiment_score (0-1), emotion_labels (array), and urgency_level (low/medium/high/critical). File: backend/app/core/sentiment_analyzer.py.'),
    ('[IC-03] Confidence Scoring Real Data: ', 'Wire confidence scoring formula to real data sources (technique execution metrics, LLM response quality scores, knowledge base retrieval relevance, conversation context length). Currently uses hardcoded formula. File: backend/app/core/confidence_scorer.py.'),
    ('[IC-09] PII NER Detection: ', 'Extend PII redaction beyond regex patterns with LLM-based named entity recognition for person names, addresses, and custom entities. Use LLM for ambiguous cases where regex confidence is low. File: backend/app/core/pii_detector.py.'),
    ('[IC-11] Hallucination Detection: ', 'Cross-check AI responses against knowledge base facts. Verify numerical claims and order details against real database data. Flag unsupported statements for human review before delivery. File: backend/app/core/hallucination_detector.py.'),
    ('[IC-12] Self-Healing Engine: ', 'Detect LLM API failures and auto-retry with fallback model via Smart Router. Monitor database connections and trigger automatic reconnection on failure. Track queue depths and alert for worker scaling. File: backend/app/core/self_healing.py.'),
]
for prefix, text in inc_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('7.3 Comprehensive Testing (12-14 hours)', level=2)
add_para(doc, 'This is the most thorough testing phase in the roadmap. It covers the full test pyramid from unit tests through integration tests to end-to-end tests, load tests, and failover tests. Every test must pass before proceeding to Week 6.')

test_items = [
    ('[TEST-01] Full existing test suite: ', 'Run pytest backend/app/tests/ with coverage reporting. Target 80%+ coverage on all files modified in Weeks 1-5. Fix any failing tests discovered during execution.'),
    ('[TEST-02] Phase verification tests: ', 'Targeted tests for every fix implemented in Weeks 1-5. Each fix must have at least one test verifying the fix works and one regression test ensuring the issue does not reoccur.'),
    ('[TEST-03] AI technique integration tests: ', 'Verify all 12 AI techniques make real LLM calls. Mock LLM responses but verify the call structure, prompt format, and response parsing logic. Test each technique\'s fallback behavior.'),
    ('[TEST-04] End-to-end pipeline test: ', 'Customer complaint through the full 19-node LangGraph pipeline with real LLM calls at each stage. Test all 4 verticals (refund, technical, billing, complaint) and all 3 tiers (Mini, PARWA, PARWA High).'),
    ('[TEST-05] Load testing: ', '100 concurrent conversations processed simultaneously. Verify no connection pool exhaustion, no memory leaks, no rate limit errors, and consistent response times. Document performance baselines.'),
    ('[TEST-06] Failover testing: ', 'Kill Redis (verify fail-closed rate limiting), kill PostgreSQL (verify graceful error handling and reconnection), simulate LLM API failures (verify Smart Router fallback to alternative providers), simulate webhook delivery failures (verify retry logic and dead letter queue).'),
    ('[TEST-07] Security regression: ', 'Verify all 93 security findings (1 CRITICAL, 22 HIGH, 39 MEDIUM, 31 LOW) are resolved. Run existing test files and create new targeted tests for previously-CRITICAL and HIGH items.'),
    ('[TEST-08] Multi-tenant isolation: ', 'Attempt cross-tenant data access from Company A to Company B via every API endpoint. Verify all attempts are blocked by both application middleware and PostgreSQL RLS. Test with direct database queries as well.'),
]
for prefix, text in test_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('7.4 Week 5 Testing Gate', level=2)
test_w5 = [
    'Unit tests for all 25 loophole detection categories (individual tests per category)',
    'Unit tests for voice intent, sentiment NLP, PII NER, hallucination detection, self-healing engine',
    'Full regression test suite combining all tests from Weeks 1-5',
    'Load test: 100 concurrent conversations with documented performance baselines',
    'Failover test: Redis, PostgreSQL, LLM API, and webhook failure scenarios all verified',
    'Multi-tenant isolation: every API endpoint tested for cross-tenant data access (all blocked)',
    'End-to-end test: all 4 verticals x 3 tiers = 12 pipeline scenarios tested',
    'Coverage target: 80%+ overall on entire modified codebase',
]
for t in test_w5:
    add_bullet(doc, t)

p = doc.add_paragraph()
r = p.add_run('GATE: All tests pass \u2192 git push \u2192 proceed to Week 6')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x0A, 0x82, 0x8A)

# ══════════════════════════════════════════
# 8. PHASE 6 (WEEK 6)
# ══════════════════════════════════════════
doc.add_heading('8. Phase 6: Deployment Readiness + Documentation (Week 6)', level=1)
add_para(doc, 'Estimated Effort: 20-25 hours across 1 week')
add_para(doc, 'Dependencies: All prior phases complete (all tests passing, zero open bugs)')
add_para(doc, 'Phase 6 is the final push before production launch. It addresses production deployment infrastructure, monitoring setup, documentation accuracy, and final quality assurance. After this phase, PARWA will be fully deployable and production-ready.')

doc.add_heading('8.1 Production Deployment (12-14 hours)', level=2)
deploy_items = [
    ('[DEPLOY-01] Docker Production Stack: ', 'Verify docker-compose.prod.yml with all services running as non-root with resource limits (CPU and memory). Test full stack startup order with proper health checks between dependent services (PostgreSQL \u2192 Redis \u2192 Backend \u2192 Frontend).'),
    ('[DEPLOY-02] SSL/TLS Configuration: ', 'TLSv1.2+ only, strong cipher suites (ECDHE-RSA-AES256-GCM-SHA384 minimum), OCSP stapling enabled, HSTS headers with preload flag. Verify certificate chain completeness. Must pass SSL Labs A+ test. File: infra/docker/nginx.conf.'),
    ('[DEPLOY-03] Database Migration Procedure: ', 'Create production migration script that runs alembic upgrade head on existing databases with live data. Test rollback procedures for every migration. Verify no data loss during migration execution. Test on a copy of production data.'),
    ('[DEPLOY-04] Backup and Recovery: ', 'Implement automated PostgreSQL backup: daily full backup, hourly WAL archiving for point-in-time recovery. Test restoration procedure on a clean environment. Create disaster recovery runbook with step-by-step instructions and recovery time objectives.'),
    ('[DEPLOY-05] Monitoring Setup: ', 'Deploy Grafana dashboards from existing configurations in monitoring/grafana_dashboards/. Configure Prometheus scraping + AlertManager rules from monitoring/ config files. Set up alerts for: API error rate > 1%, response time P95 > 2s, queue depth > 100, LLM error rate > 5%, database connection pool > 80%.'),
    ('[DEPLOY-06] AI Pipeline Monitoring: ', 'Implement custom metrics for per-technique latency, token usage per request (grouped by tier), LLM error rates by provider, Smart Router fallback frequency, FAKE Voting consensus scores, and Jarvis override rates. Export all metrics to Prometheus.'),
    ('[DEPLOY-07] Error Tracking: ', 'Implement structured error logging with Sentry (or equivalent). Configure alerts for critical error categories: unhandled exceptions, database connection failures, LLM API timeouts, authentication failures. Add request tracing with correlation IDs.'),
]
for prefix, text in deploy_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('8.2 Documentation Update (6-8 hours)', level=2)
doc_items = [
    ('[DOC-01] Architecture Document Update: ', 'Update PARWA Architecture Design Document to reflect the real AI architecture as built. Remove claims about capabilities not yet implemented. Document actual LLM integration via llm_gateway.py, real CLARA RAG with HyDE/multi-query, MAKER pipeline stages, and FAKE Voting.'),
    ('[DOC-02] Document Audit: ', 'Audit all documents in /documents/ and /docs/. Update or remove incorrect claims about AI capabilities. Ensure every document accurately describes the LLM-first with template fallback architecture rather than claiming pure LLM or pure template approaches.'),
    ('[DOC-03] Production Runbook: ', 'Create comprehensive production deployment runbook covering: environment setup (all required variables), database migration procedure, Docker deployment steps, SSL certificate installation, monitoring configuration, rollback procedures, and troubleshooting guide for common issues.'),
    ('[DOC-04] API Documentation: ', 'Regenerate OpenAPI/Swagger documentation reflecting all authentication changes, new endpoints, and modified request/response schemas. Verify documentation matches actual API behavior with automated tests.'),
    ('[DOC-05] Specification Documents: ', 'Update PARWA_AI_Technique_Framework.md, PARWA_Context_Bible.md, JARVIS_SPECIFICATION.md, and PARWA_SRS_Software_Requirements_Specification.md to accurately describe what the system actually does versus what was originally planned.'),
]
for prefix, text in doc_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('8.3 Week 6 Final Testing Gate', level=2)
test_w6 = [
    'Production deployment smoke test on staging environment',
    'SSL/TLS verification: must pass SSL Labs A+ test',
    'Database migration test on copy of production data (zero data loss)',
    'Backup and restore verification: backup \u2192 restore \u2192 verify data integrity',
    'Monitoring alerts validation: trigger each alert condition and verify notification delivery',
    'Documentation accuracy review: all docs match actual system behavior',
    'Full end-to-end production simulation: customer signup \u2192 create ticket \u2192 AI response \u2192 escalation \u2192 resolution',
    'Performance regression: compare against Week 5 baselines, no degradation',
]
for t in test_w6:
    add_bullet(doc, t)

p = doc.add_paragraph()
r = p.add_run('GATE: All tests pass \u2192 PRODUCTION DEPLOY APPROVED')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x0A, 0x82, 0x8A)

# ══════════════════════════════════════════
# 9. GAP-TO-PHASE MAPPING
# ══════════════════════════════════════════
doc.add_heading('9. Complete Gap-to-Phase Mapping', level=1)
add_para(doc, 'The following table maps every identified gap, bug, and incomplete item to its assigned phase/week, source audit category, and current status. This provides a single reference for tracking all work items across the entire roadmap.')

add_table(doc,
    ['Phase/Week', 'Gap ID', 'Category', 'Description'],
    [
        ['Week 1', 'C-07', 'Security', '.env.prod still tracked in git'],
        ['Week 1', 'C-15', 'Security', 'Empty dev pepper fallback'],
        ['Week 1', 'C-14', 'Security', 'Fernet vs AES-256-GCM decision'],
        ['Week 1', 'H-09', 'Security', 'Pricing key hardcoded fallback'],
        ['Week 1', 'AI-F01', 'AI', 'LLM client path mismatch'],
        ['Week 1', 'AI-F02', 'AI', 'Missing execute_with_llm()'],
        ['Week 1', 'GAP-FAKE', 'AI', 'FAKE Voting system for MAKER'],
        ['Week 1', 'GAP-FE1', 'Frontend', 'Dashboard: Tickets page'],
        ['Week 1', 'H-01/04/06', 'Security HIGH', 'Open redirect + CSP + IP + rate limit'],
        ['Week 1', 'H-07/08/19/20', 'Security HIGH', 'Webhook + CSRF + mock login'],
        ['Week 1', 'H-12/16/17', 'Security HIGH', 'OAuth token + email HTML + API key leak'],
        ['Week 1', 'H-13/15/21', 'Security HIGH', 'Billing role + webhook auth + rate limit'],
        ['Week 1', 'M-27/37', 'Security HIGH', 'User enum + hardcoded phone'],
        ['Week 1', 'D-02/D-03', 'Dashboard', 'Analytics mock data + email sanitization'],
        ['Week 1', 'GAP-SOCKET', 'Real-time', 'Socket.io frontend integration'],
        ['Week 2', 'AI-14/14b', 'AI RAG', 'CLARA RAG: HyDE + Multi-Query + Compression'],
        ['Week 2', 'AI-15', 'AI MAKER', 'MAKER full pipeline (6-24 LLM calls)'],
        ['Week 2', 'AI-12', 'AI DSPy', 'Real DSPy integration (no StubModule)'],
        ['Week 2', 'AI-13', 'AI Training', 'Agent Lightning real training pipeline'],
        ['Week 2', 'AI-16', 'AI Routing', '3-Tier Hybrid optimization'],
        ['Week 2', 'BUG-HT', 'Bug', 'Smart Router HealthTracker state'],
        ['Week 2', 'BUG-MV', 'Bug', 'MockVectorStore default swap'],
        ['Week 3', 'M-xx (17)', 'Security MED', '17 MEDIUM security fixes'],
        ['Week 3', 'GAP-FE2-7', 'Frontend', '6 Dashboard pages (Billing/KB/Settings/Agents/Monitor/Variants)'],
        ['Week 3', 'LG-01-12', 'LangGraph', 'All 19 LangGraph node implementations'],
        ['Week 4', 'BUG-1/2/3/4', 'Critical Bug', 'Migration FK + UUID + route prefix + Jarvis'],
        ['Week 4', 'GAP-RLS', 'Security', 'Row Level Security (93 tables)'],
        ['Week 4', 'L-xx + INF', 'Infra', '13 LOW findings + Docker + env vars'],
        ['Week 5', 'GAP-LOOPHOLE', 'Safety', '25 Loophole Solutions module'],
        ['Week 5', 'IC-01-12', 'Incomplete', 'Voice + Sentiment + PII + Hallucination + Self-Healing'],
        ['Week 5', 'TEST-01-08', 'Testing', 'Full regression + load + failover + security tests'],
        ['Week 6', 'DEPLOY-01-07', 'Deployment', 'Docker + SSL + backup + monitoring'],
        ['Week 6', 'DOC-01-05', 'Documentation', 'Architecture + marketing + runbook + API docs'],
    ],
    col_widths=[0.9, 1.2, 1.0, 2.9]
)

# ══════════════════════════════════════════
# 10. GIT WORKFLOW AND TESTING PROTOCOL
# ══════════════════════════════════════════
doc.add_heading('10. Git Workflow and Testing Protocol', level=1)
add_para(doc, 'This section defines the mandatory git workflow and testing protocol that must be followed throughout all 6 weeks of this roadmap. Adherence to this protocol is non-negotiable and ensures code quality, prevents regressions, and maintains a deployable codebase at all times.')

doc.add_heading('10.1 Strict Git Workflow', level=2)
add_para(doc, 'Every development session must follow this exact git workflow. No code may be pushed without completing the full workflow. This applies to all developers working on the PARWA codebase.')

git_items = [
    'At the start of every development session: git pull origin main to get the latest changes',
    'Create a feature branch for the current week: git checkout -b week-N-phase-name (e.g., week-1-security-high)',
    'Complete the assigned work items for the current week/phase',
    'Run the full test suite for the current week\'s testing gate',
    'Fix any failing tests before proceeding',
    'Only after ALL tests pass: git add, git commit, git push origin week-N-phase-name',
    'Create a Pull Request from the feature branch to main',
    'After PR approval (or self-review for solo developer): merge to main',
    'Pull latest main before starting the next week\'s work: git pull origin main',
]
for item in git_items:
    add_bullet(doc, item)

doc.add_heading('10.2 Testing Protocol Per Week', level=2)
add_para(doc, 'Each week MUST pass through the following testing gate before any code is pushed to the repository. This is a mandatory quality checkpoint that prevents regressions and ensures the system remains stable throughout the development process. The testing gate consists of five layers:')

test_protocol = [
    ('1. Unit Tests: ', 'Run pytest for all new and modified code. Every function, class, and component must have corresponding unit tests. Target: minimum 80% code coverage on all modified files. Use pytest-cov for coverage reporting.'),
    ('2. Integration Tests: ', 'Test API endpoints end-to-end (request to response), database operations (CRUD with real database), cross-service communication (backend to frontend, backend to LLM providers, backend to external APIs like Paddle/Twilio/Brevo).'),
    ('3. Security Regression: ', 'Run the security regression test suite after every change. Verify that no new CRITICAL or HIGH vulnerabilities have been introduced. Use automated security scanning tools where available.'),
    ('4. Performance Baseline: ', 'Compare current performance against the baseline established in the previous week. No degradation allowed in API response times, database query performance, or memory usage without explicit justification.'),
    ('5. Documentation: ', 'All new code must be documented with docstrings. API documentation (OpenAPI/Swagger) must be updated for any new or modified endpoints. Architecture documentation must reflect any structural changes.'),
]
for prefix, text in test_protocol:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('10.3 Parallel Work Handling', level=2)
add_para(doc, 'Since development is happening on multiple features in parallel (the developer is working on multiple things simultaneously), the following rules apply to prevent merge conflicts and maintain code quality:')

parallel_items = [
    'Always pull the latest code from main before starting any new work session',
    'Resolve merge conflicts immediately when they occur, never leave them pending',
    'Run the full test suite after any merge operation, even for simple conflict resolution',
    'Never push code to main without passing all tests in the weekly testing gate',
    'If working on multiple branches simultaneously, rebase frequently against main',
    'Communicate branch status clearly: which branches are in progress, which are ready for merge',
]
for item in parallel_items:
    add_bullet(doc, item)

# ══════════════════════════════════════════
# 11. EFFORT SUMMARY
# ══════════════════════════════════════════
doc.add_heading('11. Effort Summary and Risk Assessment', level=1)

add_table(doc,
    ['Phase/Week', 'Focus Area', 'Hours', 'Primary Risk'],
    [
        ['Week 1', 'Security Carry-Forward + HIGH + Socket.io', '35-40', 'FAKE Voting prompt quality; Socket.io auth'],
        ['Week 2', 'AI Frameworks + MAKER + RAG', '35-40', 'DSPy complexity; training API differences'],
        ['Week 3', 'MEDIUM Security + Dashboards + LangGraph', '40-45', 'LangGraph node interdependencies'],
        ['Week 4', 'RLS + Critical Bugs + Infrastructure', '35-40', 'Migration breakage; RLS performance'],
        ['Week 5', 'Loopholes + Incomplete + Full Testing', '40-45', 'Loophole detection accuracy; test coverage'],
        ['Week 6', 'Deployment + Documentation', '20-25', 'SSL config; documentation accuracy'],
        ['TOTAL', '6 Weeks', '205-235', '\u2014'],
    ],
    col_widths=[0.9, 2.8, 0.8, 2.5]
)

# ══════════════════════════════════════════
# 12. SUCCESS METRICS
# ══════════════════════════════════════════
doc.add_heading('12. Success Metrics', level=1)
add_para(doc, 'The roadmap succeeds when the following metrics are achieved at the end of Week 6:')

metrics = [
    'Zero CRITICAL vulnerabilities remain (currently 1: C-07, resolved in Week 1)',
    'Zero HIGH vulnerabilities remain unaddressed (all 22 resolved by Week 1)',
    'Minimum 80% of MEDIUM findings resolved (31+ of 39 by Week 3)',
    'All 17 LOW findings resolved (Week 4)',
    'All 12 AI techniques confirmed making real LLM API calls (verified by Day 3 cross-check)',
    'CLARA RAG performs HyDE generation, multi-query retrieval, and contextual compression (Week 2)',
    'MAKER Framework executes 6+ LLM calls per query with FAKE Voting consensus (Week 1-2)',
    'Agent Lightning collects mistakes, triggers training, and deploys models (Week 2)',
    'All 19 LangGraph nodes perform real processing with no pass-through stubs (Week 3)',
    'All 7 dashboard pages display real data from backend APIs (Week 1-3)',
    'Socket.io provides real-time updates across all dashboard pages (Week 1)',
    'PostgreSQL RLS active on all 93 tenant-scoped tables (Week 4)',
    '25 Loophole Solutions module detects and counters all defined categories (Week 5)',
    'Fresh database deployment succeeds with alembic upgrade head (Week 4)',
    'Full test suite passes with no regressions, 80%+ coverage on all modified files (Week 5)',
    '100 concurrent conversations handled stably with no degradation (Week 5)',
    'All documentation accurately describes real system capabilities (Week 6)',
    'Production deployment passes staging smoke test and SSL Labs A+ (Week 6)',
]
for m in metrics:
    add_bullet(doc, m)

# ── Save ──
output_path = '/home/z/my-project/download/PARWA_Production_Readiness_Roadmap.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')

import os
size = os.path.getsize(output_path)
print(f'File size: {size:,} bytes ({size/1024:.1f} KB)')
