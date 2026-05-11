#!/usr/bin/env python3
"""Generate PARWA Production Readiness Roadmap - Verified 3-4 Week Plan"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hf = hs.font
    hf.name = 'Calibri'
    hf.bold = True
    if level == 1:
        hf.size = Pt(18); hf.color.rgb = RGBColor(0x0A, 0x16, 0x28)
    elif level == 2:
        hf.size = Pt(14); hf.color.rgb = RGBColor(0x0A, 0x16, 0x28)
    else:
        hf.size = Pt(12); hf.color.rgb = RGBColor(0x1A, 0x2B, 0x40)

# ── Helpers ──
NB = { 'style': 'none', 'sz': 0, 'color': 'FFFFFF', 'space': 0 }

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]; cell.text = h
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F2027"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]; cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs: run.font.size = Pt(9)
        if r_idx % 2 == 0:
            for c_idx in range(len(headers)):
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F8FC"/>')
                row.cells[c_idx]._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths): row.cells[i].width = Inches(w)
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(10)
        r = p.add_run(text); r.font.size = Pt(10)
    else:
        r = p.add_run(text); r.font.size = Pt(10)
    return p

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    return p

def gate_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x0A, 0x82, 0x8A)
    return p

# ══════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════
for _ in range(6): doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PARWA'); r.bold = True; r.font.size = Pt(42); r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Production Readiness Roadmap'); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Code-Verified Execution Plan \u2014 4 Weeks to Production'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x50, 0x60, 0x70)

for _ in range(3): doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('All 32 items cross-checked against actual source code'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x60, 0x70, 0x80)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('23 confirmed missing | 5 already done (false positives) | 4 partial (nginx config drift)'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x60, 0x70, 0x80)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('May 2025'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x60, 0x70, 0x80)

doc.add_page_break()

# ══════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
for item in [
    '1. Executive Summary',
    '2. Cross-Check Audit Results',
    '3. Week 1: AI Engine + Security Hardening',
    '4. Week 2: AI Upgrades + Dashboard Pages',
    '5. Week 3: Database + Infrastructure + Loopholes',
    '6. Week 4: Deployment + Testing + Documentation',
    '7. Complete Gap-to-Week Mapping',
    '8. Git Workflow and Testing Protocol',
    '9. Effort Summary',
    '10. Success Metrics',
]:
    add_bullet(doc, item)

add_para(doc, '(Right-click on the table of contents above and select "Update Field" to refresh page numbers.)', size=9)
doc.add_page_break()

# ══════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

add_para(doc, 'This roadmap takes PARWA from its current state (~82% built) to full production readiness in 4 weeks. Unlike the previous 6-week plan which included 16 false-positive items (already built), this version is based on a rigorous cross-check of every claimed gap against the actual source code across 2,200+ Python files, 16MB backend, and ~145 frontend files.')

add_para(doc, 'The cross-check verified 32 items that were claimed as missing. Results: 5 items were false positives (already fully implemented), 4 items were partial (mostly nginx config drift between two config files), and 23 items are genuinely missing and need to be built. This reduces the estimated effort from 6 weeks to 4 focused weeks.')

add_para(doc, 'The work is organized into 4 weekly phases. Each phase concludes with a mandatory testing gate: unit tests, integration tests, and security regression. Code may only be pushed after all tests pass. This strict protocol prevents regressions and ensures the system remains deployable throughout development.')

add_para(doc, 'Key finding: The most impactful gap is the nginx configuration drift. The file nginx/nginx.conf has production-grade security (OCSP stapling, full HSTS, 8 cipher suites), but infra/docker/nginx-default.conf (the actual Docker deployment config) is significantly weaker, missing OCSP stapling, HSTS includeSubDomains/preload flags, and having only 2 cipher suites instead of 8. This single fix (syncing docker config with the main config) resolves 3 partial items at once.')

# ══════════════════════════════════════════
# 2. CROSS-CHECK AUDIT RESULTS
# ══════════════════════════════════════════
doc.add_heading('2. Cross-Check Audit Results', level=1)

add_para(doc, 'Every item from the claimed "32 genuinely missing" list was verified against the actual codebase. The results are summarized below with detailed evidence for each verdict.')

doc.add_heading('2.1 Items Already Done (False Positives \u2014 No Work Needed)', level=2)
add_para(doc, 'These 5 items were claimed as missing but are fully implemented in the codebase. No work is required.')

add_table(doc,
    ['Item', 'Claim', 'Actual State (Verified)', 'Evidence'],
    [
        ['UUID Type Inconsistency', '2 models use UUID vs String(36)', 'BOTH files use UUID(as_uuid=True) consistently', 'email_delivery_event.py + outbound_email.py'],
        ['.env.production.template', 'Missing', 'EXISTS as .env.prod.example (85 lines)', 'Root directory'],
        ['M-04: email-validator', 'Simple @ validation', 'email-validator>=2.1.0 in requirements.txt', 'requirements.txt line 34 + verification.py'],
        ['M-06: API Key 401', 'Pass-through on missing header', 'Returns 401 for invalid API keys', 'api_key_auth.py lines 77-86'],
        ['M-16: Twilio Signature', 'No SMS callback verification', 'Full HMAC signature verification', 'sms_channel.py lines 747-798'],
    ],
    col_widths=[1.2, 1.5, 1.8, 1.5]
)

doc.add_heading('2.2 Items Partially Done (Quick Fixes)', level=2)
add_para(doc, 'These 4 items have some implementation but are incomplete. Most are caused by nginx config drift between two files.')

add_table(doc,
    ['Item', 'What Exists', 'What\'s Missing', 'Fix'],
    [
        ['OCSP Stapling', 'In nginx/nginx.conf (lines 100-101)', 'NOT in infra/docker/nginx-default.conf', 'Sync docker config with main config'],
        ['HSTS includeSubDomains', 'In nginx/nginx.conf (line 106)', 'NOT in docker config (only max-age)', 'Add includeSubDomains; preload to docker'],
        ['TLS Cipher Suites', '8 ciphers in nginx/nginx.conf', 'Only 2 in docker config', 'Copy full cipher list to docker config'],
        ['JWT Token Blacklist', 'jti claim generated in tokens', 'No is_token_revoked() function', 'Add Redis blacklist check + cleanup'],
    ],
    col_widths=[1.2, 1.5, 1.5, 1.8]
)

add_para(doc, 'CRITICAL FINDING: The nginx config drift is the single most impactful discovery. The infra/docker/nginx-default.conf (which runs in production via Docker) has significantly weaker security than nginx/nginx.conf. This must be fixed by syncing the Docker config to match the main config, which resolves 3 of the 4 partial items in one change.')

doc.add_heading('2.3 Dashboard Pages Status', level=2)
add_para(doc, 'All 7 dashboard pages exist. 2 are fully built (Tickets at 1,362 lines, Variants at 722 lines). 5 are stubs at `src/app/dashboard/` that need real implementation with backend API integration. The stubs range from 15 to 52 lines with "Coming Soon" placeholders.')

add_table(doc,
    ['Page', 'Status', 'Lines', 'What\'s Needed'],
    [
        ['Tickets', 'REAL', '1,362', 'Already built with real data'],
        ['Variants', 'REAL', '722', 'Already built with real data'],
        ['Billing', 'STUB', '52', 'Plans, invoices, payment methods, upgrade/downgrade'],
        ['Knowledge Base', 'STUB', '52', 'Upload, search, document list, reindex'],
        ['Settings', 'STUB', '15', 'Profile, API keys, team, notifications, security'],
        ['AI Agents', 'STUB', '15', 'Agent list, metrics, capability matrix'],
        ['Monitoring', 'STUB', '52', 'System health, AI metrics, alerts, SLA'],
    ],
    col_widths=[1.2, 0.6, 0.5, 3.7]
)

# ══════════════════════════════════════════
# 3. WEEK 1
# ══════════════════════════════════════════
doc.add_heading('3. Week 1: AI Engine + Security Hardening', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours | Focus: CLARA RAG, Training Pipeline, FAKE Voting, Security HIGH')
add_para(doc, 'Week 1 delivers the most impactful AI features and closes the highest-priority security gaps. CLARA RAG gets its advanced retrieval capabilities (HyDE + Multi-Query). The Agent Lightning training pipeline goes from hardcoded stubs to real implementation. FAKE Voting adds multi-evaluator consensus to the MAKER Framework. Critical security fixes address webhook replay, rate limiting, and JWT hardening.')

doc.add_heading('3.1 CLARA RAG Advanced Retrieval (8-10 hours)', level=2)
add_para(doc, 'CLARA RAG currently has a quality gate pipeline (clara_quality_gate.py, 704 lines) and RAG re-ranking (rag_reranking.py, 1,596 lines using BM25/keyword scoring), but lacks LLM-powered retrieval features. Context compression already exists (context_compression.py, 749 lines) with 5 strategies. The missing pieces are HyDE generation and multi-query retrieval.')

add_para(doc, '[MISSING-1] HyDE (Hypothetical Document Embedding): When a customer query arrives, use the LLM to generate a hypothetical answer document via llm_gateway.py. Embed this hypothetical answer and use it as the primary retrieval vector. This dramatically improves retrieval quality because the embedding space is optimized for document embeddings, not query embeddings. File: backend/app/core/rag/clara_rag.py (new file).')

add_para(doc, '[MISSING-2] Multi-Query Retrieval: Generate 3 alternative phrasings of the original query using the LLM, then retrieve documents for all 4 queries (original + 3 alternatives). Merge, deduplicate, and rank by aggregate relevance score. Handles vocabulary mismatch between customer language and knowledge base terminology. File: backend/app/core/rag/multi_query.py (new file).')

add_para(doc, '[EXISTING] Context Compression: Already built at context_compression.py (749 lines) with 5 strategies. No work needed here. RAG Re-ranking exists at rag_reranking.py (1,596 lines) but uses BM25/keyword scoring. Consider LLM-based re-ranking as a stretch goal.')

doc.add_heading('3.2 Agent Lightning Training Pipeline (6-8 hours)', level=2)
add_para(doc, '[MISSING-5] The training pipeline currently consists of 3 Celery task stubs (training_tasks.py, 140 lines) that return hardcoded zeros: samples_count: 0, training_triggered: False, training_job_id: None. Database models exist (TrainingDataset, TrainingCheckpoint, AgentMistake, AgentPerformance, TrainingRun) but no logic connects them.')

add_para(doc, 'Build the real training pipeline: (1) Mistake collection: query AgentMistake table for FAKE Voting red-flags and Jarvis overrides since last training run. (2) Data preparation: convert mistakes into LLaMA-3-8B fine-tuning format (instruction/input/output triples) via Unsloth-compatible dataset. (3) Training trigger: when mistakes exceed configurable threshold (default 50), automatically prepare dataset and trigger training. (4) Google Colab integration: generate a download link for the prepared dataset and a pre-configured Colab notebook URL. (5) Model deployment: endpoint to upload the fine-tuned model weights and activate them for a specific agent/technique. (6) A/B testing: split traffic between default and fine-tuned model, compare resolution rate and customer satisfaction.')

doc.add_heading('3.3 FAKE Voting + Security HIGH Fixes (12-14 hours)', level=2)
add_para(doc, '[GAP-FAKE] Implement FAKE Voting in MAKER validator (11_maker_validator.py, 684 lines): generate 3-5 candidate responses per query, score each using 4+ evaluation dimensions (fluency, relevance, safety, brand voice), apply Red-Flagging for hallucinations and policy violations, produce weighted consensus score.')

add_para(doc, '[MISSING-12] Rate Limiting Fail-Closed: Change rate_limit.py line 91 from fail-open to fail-closed. When Redis is down, block all requests (return 429) instead of allowing all. This is a critical security hardening \u2014 fail-open rate limiting means an attacker who can knock out Redis gets unlimited access.')

add_para(doc, '[MISSING-13] Webhook Replay Protection: Add timestamp freshness validation to billing_webhooks.py. Reject webhook payloads where occurred_at is older than 5 minutes. Apply to all webhook handlers (Paddle, Twilio, Brevo, Shopify).')

add_para(doc, '[MISSING-6] JWT Token Blacklist: jti claim is generated in tokens (auth.py line 98-99) and get_token_jti() helper exists (line 149), but no is_token_revoked() function exists. Implement: Redis SET for blacklisted jtis with TTL matching token expiry, check on every authenticated request, Celery beat cleanup task for expired entries.')

add_para(doc, '[MISSING-16] Shared IP Extraction: Create get_client_ip() in backend/app/core/utils.py. Currently 3 duplicate implementations exist: rate_limit.py (line 121, method), ip_allowlist.py (line 152, method with ASGI scope), request_logger.py (line 73, standalone function). Consolidate into one shared utility.')

doc.add_heading('3.4 Week 1 Testing Gate', level=2)
for t in [
    'Unit tests for CLARA RAG: HyDE generation, multi-query retrieval, dedup, ranking',
    'Unit tests for training pipeline: mistake collection, data prep, threshold trigger, A/B framework',
    'Unit tests for FAKE Voting: candidate generation, scoring, red-flagging, consensus',
    'Unit tests for rate limit fail-closed (Redis down = 429), webhook timestamp (reject stale)',
    'Integration test for JWT blacklist: revoke token, verify subsequent requests rejected',
    'Integration test for shared IP utility: verify consistent behavior across all 3 consumers',
    'Coverage target: 80%+ on all modified files',
]:
    add_bullet(doc, t)
gate_para(doc, 'GATE: All tests pass \u2192 git push \u2192 proceed to Week 2')

# ══════════════════════════════════════════
# 4. WEEK 2
# ══════════════════════════════════════════
doc.add_heading('4. Week 2: AI Upgrades + Dashboard Pages', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours | Focus: LLM upgrades to Sentiment/PII/Voice, 5 Dashboard pages, Socket.io')
add_para(doc, 'Week 2 upgrades three AI components from regex/keyword-based to LLM-powered, builds all 5 remaining dashboard stub pages with real backend API integration, and adds the Socket.io frontend client for real-time updates.')

doc.add_heading('4.1 LLM-Powered AI Upgrades (10-12 hours)', level=2)

add_para(doc, '[MISSING-10] LLM-Based Sentiment Analysis: sentiment_engine.py (855 lines) currently uses static word lists (FRUSTRATION_STRONG, POSITIVE_WORDS, EMOTION_LEXICON) and pattern matching (ALL CAPS, exclamation density). Upgrade: add LLM fallback that produces structured output (sentiment_score: 0-1, emotion_labels: [], urgency_level: low/medium/high/critical) via llm_gateway.py. Keep lexicon as fast path for simple cases, use LLM for ambiguous/complex sentiment. This is not a rewrite \u2014 add an LLM enrichment layer on top of the existing engine.')

add_para(doc, '[MISSING-11] LLM-Based PII NER: pii_redaction_engine.py (1,155 lines) uses 16 compiled regex patterns. Add LLM-based named entity recognition for ambiguous cases where regex confidence is low (person names, addresses, custom entities). Use LLM as a second pass after regex: regex catches known patterns, LLM catches context-dependent PII that regex misses. This hybrid approach keeps the speed of regex while adding LLM accuracy for edge cases.')

add_para(doc, '[MISSING-9] Voice Server LLM Intent: parwa_voice_server.py _detect_intent() (line 444) uses pure keyword matching with 9 static intent categories. Connect to Smart Router for LLM-based intent classification. Keep keyword matching as fast path for obvious intents (greeting, goodbye), use LLM via Smart Router for ambiguous or complex intents. The comment at line 447-448 already says "in production, this routes through Parwa\'s full variant pipeline" \u2014 make it true.')

doc.add_heading('4.2 Dashboard Pages (16-18 hours)', level=2)
add_para(doc, 'Build all 5 stub pages with real backend API integration, real-time Socket.io updates, proper loading states, error handling, and responsive design. Each page must connect to actual backend endpoints and display live data.')

dash_items = [
    ('[Billing] ', 'Subscription overview with current plan, renewal date, usage meters. Invoice list with PDF download. Payment method management (add/remove/default). Upgrade/downgrade via Paddle integration. File: src/app/dashboard/billing/page.tsx (currently 52 lines stub).'),
    ('[Knowledge Base] ', 'Document upload with drag-and-drop and progress. Document list with status (indexed/indexing/failed). Full-text search. Manual reindex trigger. Document deletion. File: src/app/dashboard/knowledge/page.tsx (currently 52 lines stub).'),
    ('[Settings] ', 'Tabbed interface: Company Profile (name, logo, industry), API Keys (generate, view, revoke), Team Management (invite, roles, remove), Notification Preferences (toggles per category), Security (MFA, sessions, IP allowlist). File: src/app/dashboard/settings/page.tsx (currently 15 lines).'),
    ('[AI Agents] ', 'Agent list with status (active/training/idle). Performance metrics (resolution rate, response time, satisfaction). Capability matrix per variant. Training history. Configuration panel. File: src/app/dashboard/agents/page.tsx (currently 15 lines).'),
    ('[Monitoring] ', 'System health (API latency P50/P95, error rates, queue depths). AI engine metrics (token usage, confidence distribution, technique counts). Alert feed. SLA compliance. File: src/app/dashboard/monitoring/page.tsx (currently 52 lines stub).'),
]
for prefix, text in dash_items:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('4.3 Socket.io Frontend Client (6-8 hours)', level=2)
add_para(doc, '[MISSING-3] The backend has a full Socket.io server (socketio.py, 443 lines) with rooms, JWT auth, and event buffering. The frontend has ZERO Socket.io code \u2014 socket.io-client is not even in package.json. Build: (1) Install socket.io-client. (2) Create useSocket React hook managing connection, auth token passing, reconnection with exponential backoff, and event subscriptions. (3) Real-time ticket status updates on the Tickets page. (4) Live chat message streaming on the customer-facing widget. (5) Notification push (new ticket, escalation, SLA warnings) across all dashboard pages. (6) Event replay for missed messages during disconnection.')

doc.add_heading('4.4 Week 2 Testing Gate', level=2)
for t in [
    'Unit tests for sentiment LLM upgrade: structured output parsing, fallback behavior',
    'Unit tests for PII NER LLM: low-confidence regex triggers LLM, high-confidence skips LLM',
    'Unit tests for voice intent: keyword fast path + LLM fallback + Smart Router integration',
    'Unit tests for all 5 dashboard pages: rendering, API calls, filters, error states',
    'Unit tests for useSocket hook: connect, auth, events, reconnect, missed message replay',
    'Integration test: Socket.io connects to backend, receives real-time ticket updates',
    'Integration test: all 5 dashboard pages load real data from backend APIs',
    'Coverage target: 80%+ on all modified files',
]:
    add_bullet(doc, t)
gate_para(doc, 'GATE: All tests pass \u2192 git push \u2192 proceed to Week 3')

# ══════════════════════════════════════════
# 5. WEEK 3
# ══════════════════════════════════════════
doc.add_heading('5. Week 3: Database + Infrastructure + Loopholes', level=1)
add_para(doc, 'Estimated Effort: 35-40 hours | Focus: RLS, Critical Bugs, Nginx hardening, Loophole Engine, Backups')
add_para(doc, 'Week 3 is the infrastructure hardening phase. PostgreSQL RLS provides defense-in-depth for multi-tenant data isolation. Critical database bugs get fixed. Nginx config drift is resolved. The 25 Loophole Solutions safety engine is built. Database backup/recovery procedures are established.')

doc.add_heading('5.1 PostgreSQL Row-Level Security (10-12 hours)', level=2)
add_para(doc, '[MISSING-2] Zero RLS policies exist on any of the 93 tenant-scoped tables. Application-level tenant isolation via SQLAlchemy middleware exists but provides no protection against SQL injection, ORM bypass, or direct database access. Build: (1) Alembic migration enabling RLS on all tables with company_id. (2) CREATE POLICY for each table: restrict SELECT, UPDATE, INSERT, DELETE to rows matching current tenant. (3) PostgreSQL function extracting company_id from JWT passed as session variable (SET app.current_tenant_id). (4) Modify connection factory to set tenant context on every checkout. (5) RLS bypass role for migrations/admin operations. (6) Test: direct SQL queries without tenant context return zero rows.')

doc.add_heading('5.2 Critical Database Bug Fixes (4-6 hours)', level=2)
add_para(doc, '[MISSING-14] FK Mismatch Migration: Migrations 003 and 005 reference sessions.id as FK but the table was renamed to tickets (BL01) without updating migrations. This causes alembic upgrade head to fail on fresh databases. Create migration 020 that handles both cases. File: database/alembic/versions/020_fix_session_ticket_fk.py.')

add_para(doc, '[NOTE: UUID Inconsistency] Cross-check confirmed this is a FALSE POSITIVE. Both email_delivery_event.py and outbound_email.py use UUID(as_uuid=True) consistently. No migration needed.')

doc.add_heading('5.3 Nginx Config Sync + Infrastructure (6-8 hours)', level=2)
add_para(doc, 'CRITICAL: Sync infra/docker/nginx-default.conf (the production Docker config) with nginx/nginx.conf (which has the correct security settings). This single change resolves 3 partial items at once:')

for item in [
    '[MISSING-21] OCSP Stapling: Add ssl_stapling on; ssl_stapling_verify on; with resolver to docker config',
    '[MISSING-22] HSTS: Add includeSubDomains; preload to the Strict-Transport-Security header',
    '[MISSING-24] TLS Ciphers: Expand from 2 to 8 ciphers (add AES-256-GCM, CHACHA20-POLY1305, DHE-RSA)',
    '[MISSING-17] Nginx Non-Root: Add USER nginx directive to infra/docker/nginx.Dockerfile (chown already exists at line 29)',
    '[MISSING-23] AlertManager Receivers: Add email/Slack receiver as backup to the single webhook receiver in alertmanager.yml',
    '[MISSING-18] PRICING_SIGNING_KEY Guard: Move from raw os.environ.get() in pricing.py line 663 to config.py with production ValueError guard',
    '[MISSING-19] ENVIRONMENT Enum: Add validator ensuring ENVIRONMENT is one of [development, staging, test, production] in config.py',
]:
    add_bullet(doc, item)

doc.add_heading('5.4 Database Backup & Recovery (4-6 hours)', level=2)
add_para(doc, '[MISSING-7] Zero backup scripts, no WAL archiving, no recovery procedures exist. PostgreSQL runs in Docker with local volumes only. Build: (1) pg_dump script for daily full backups with timestamp naming. (2) WAL archiving configuration for point-in-time recovery. (3) Automated backup cron job running inside Docker. (4) Restore procedure script tested on a clean environment. (5) Disaster recovery runbook with RTO/RPO targets. (6) Backup verification: automated restore test on schedule.')

doc.add_heading('5.5 25 Loophole Solutions Engine (8-10 hours)', level=2)
add_para(doc, '[MISSING-4] No centralized loophole detection system exists. Loophole patterns are enumerated in Building Codes documentation and 15+ test files reference them, but there is no loophole_engine.py or LoopholeRegistry implementation. Build: (1) LoopholeRegistry dataclass mapping each of 25 loophole categories to severity, detection pattern, and countermeasure. (2) LoopholeDetectionEngine in backend/app/core/loophole_engine.py that scans AI responses using both rule-based checks and LLM-based analysis. (3) Integration into LangGraph Node 14 (Guardrails) as a mandatory pre-delivery check. Categories include: hallucination of policies, PII leakage, unauthorized data access, emotional manipulation, biased responses, off-topic responses, escalation failures, brand voice violations, regulatory non-compliance, and more.')

doc.add_heading('5.6 Security MEDIUM Fixes (Remaining) (4-6 hours)', level=2)
add_para(doc, 'These MEDIUM items were verified as genuinely missing and need implementation:')

for item in [
    '[M-01] Remove user_role from AuthorizationError details \u2014 return generic "access denied"',
    '[M-08] Add explicit auth dependency to events API endpoint',
    '[M-13] Replace setattr-based user update with explicit field whitelist in admin.py',
    '[M-14] Add Pydantic request models to ai_engine endpoints accepting raw dict bodies',
    '[M-17] Replace str(e) with generic errors in knowledge base API responses',
    '[M-19] Fix visitor token verification that passes on exception instead of rejecting',
    '[M-23] Fix MCP server CORS that falls back to ["*"] on exception',
    '[M-26] Add security headers (X-Content-Type-Options, X-Frame-Options, HSTS, CSP) to Next.js middleware',
    '[M-28] Sanitize email content before passing to Brevo API',
    '[M-32] Add max payload size validation (1MB) to Celery task definitions',
    '[M-33] Escape % and _ in search queries before ILIKE to prevent SQL wildcard injection',
    '[M-35] Add role check to notification send endpoint',
]:
    add_bullet(doc, item)

doc.add_heading('5.7 Week 3 Testing Gate', level=2)
for t in [
    'Unit tests for RLS: direct SQL returns no rows without tenant context, admin bypass works',
    'Unit tests for FK migration: alembic upgrade head succeeds on fresh + existing databases',
    'Unit tests for nginx config: OCSP, HSTS, ciphers all present in docker config',
    'Unit tests for all 12 MEDIUM security fixes (individual tests)',
    'Unit tests for all 25 loophole detection categories',
    'Integration test: backup script runs, restore succeeds on clean database, data integrity verified',
    'Integration test: multi-tenant isolation via BOTH RLS and middleware (dual-layer)',
    'Coverage target: 80%+ on all modified files',
]:
    add_bullet(doc, t)
gate_para(doc, 'GATE: All tests pass \u2192 git push \u2192 proceed to Week 4')

# ══════════════════════════════════════════
# 6. WEEK 4
# ══════════════════════════════════════════
doc.add_heading('6. Week 4: Deployment + Testing + Documentation', level=1)
add_para(doc, 'Estimated Effort: 30-35 hours | Focus: JWT RS256, Full test suite, Deployment, Documentation')
add_para(doc, 'Week 4 is the final quality and deployment phase. JWT migrates from HS256 to RS256 for asymmetric verification. The most comprehensive test suite in the project runs: unit, integration, load, failover, and security regression. Production deployment is validated. Documentation is updated to match reality.')

doc.add_heading('6.1 JWT RS256 Migration (6-8 hours)', level=2)
add_para(doc, '[MISSING-6] Currently HS256 only (auth.py line 27: JWT_ALGORITHM = "HS256"). Migrate to RS256 (asymmetric) allowing frontend token verification without shared signing key. Build: (1) Generate RSA-2048 key pair (openssl genrsa). (2) Configure backend to sign with RS256 private key. (3) Configure frontend to verify with RS256 public key (no shared secret needed). (4) Support key rotation: multiple active keys with versioned metadata in Redis. (5) Graceful rotation without invalidating existing sessions. (6) Token blacklist (built in Week 1) continues to work with RS256 tokens.')

doc.add_heading('6.2 Comprehensive Testing (12-14 hours)', level=2)
add_para(doc, 'The most thorough testing phase. Covers the full test pyramid from unit through end-to-end, load, and failover. Every test must pass before production deploy approval.')

for prefix, text in [
    ('[TEST-01] Full Regression: ', 'Run pytest backend/app/tests/ with 80%+ coverage on all modified files. Fix any failing tests.'),
    ('[TEST-02] Phase Verification: ', 'Targeted tests for every fix in Weeks 1-3. Each fix has a test verifying it works + a regression test.'),
    ('[TEST-03] AI Integration: ', 'Verify all 12 techniques make real LLM calls. Mock responses but verify call structure, prompt format, response parsing.'),
    ('[TEST-04] End-to-End Pipeline: ', 'Customer query through full 19-node LangGraph pipeline with real LLM calls. Test all 4 verticals (refund, technical, billing, complaint) and 3 tiers.'),
    ('[TEST-05] Load Testing: ', '100 concurrent conversations. No connection pool exhaustion, no memory leaks, no rate limit errors. Document performance baselines.'),
    ('[TEST-06] Failover Testing: ', 'Kill Redis (verify fail-closed rate limiting), kill PostgreSQL (verify error handling), simulate LLM API failures (verify Smart Router fallback), simulate webhook failures (verify retry).'),
    ('[TEST-07] Security Regression: ', 'Verify all 93 security findings resolved. Run existing test files. Create new tests for previously-CRITICAL and HIGH items.'),
    ('[TEST-08] Multi-Tenant Isolation: ', 'Attempt cross-tenant data access via every API endpoint. Verify blocked by both middleware and RLS.'),
]:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('6.3 Production Deployment (6-8 hours)', level=2)
for prefix, text in [
    ('[DEPLOY-01] Docker Stack: ', 'Verify docker-compose.prod.yml: all services non-root, resource limits, health checks, startup order.'),
    ('[DEPLOY-02] SSL/TLS: ', 'Verify certificate chain. Must pass SSL Labs A+ test (now possible with synced nginx config).'),
    ('[DEPLOY-03] DB Migration: ', 'Test alembic upgrade head on copy of production data. Verify no data loss. Test rollback.'),
    ('[DEPLOY-04] Monitoring: ', 'Deploy Grafana + Prometheus + AlertManager. Configure alerts for: API error >1%, P95 >2s, queue >100, LLM error >5%.'),
    ('[DEPLOY-05] Error Tracking: ', 'Sentry integration with structured logging, request tracing, correlation IDs.'),
]:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('6.4 Documentation (4-6 hours)', level=2)
for prefix, text in [
    ('[DOC-01] Architecture Doc: ', 'Update to reflect real AI architecture. Remove unimplemented capability claims. Document actual LLM integration via llm_gateway.py.'),
    ('[DOC-02] Document Audit: ', 'Audit all docs in /documents/ and /docs/. Fix incorrect claims about AI capabilities.'),
    ('[DOC-03] Runbook: ', 'Production deployment runbook: env setup, DB migration, Docker deploy, SSL, monitoring, rollback, troubleshooting.'),
    ('[DOC-04] API Docs: ', 'Regenerate OpenAPI/Swagger reflecting all auth changes and new endpoints.'),
]:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('6.5 Week 4 Final Testing Gate', level=2)
for t in [
    'Production smoke test on staging environment',
    'SSL Labs A+ test (nginx config now has full ciphers + OCSP + HSTS)',
    'Database migration on production data copy (zero data loss)',
    'Backup restore verification (data integrity confirmed)',
    'Full end-to-end production simulation: signup > ticket > AI response > escalation > resolution',
    'Performance regression: compare against Week 3 baselines',
    '100 concurrent conversations stable with documented baselines',
]:
    add_bullet(doc, t)
gate_para(doc, 'GATE: All tests pass \u2192 PRODUCTION DEPLOY APPROVED')

# ══════════════════════════════════════════
# 7. GAP MAPPING
# ══════════════════════════════════════════
doc.add_heading('7. Complete Gap-to-Week Mapping', level=1)

add_table(doc,
    ['Week', 'ID', 'Category', 'Item', 'Verdict'],
    [
        ['W1', '1', 'AI RAG', 'CLARA RAG: HyDE + Multi-Query', 'MISSING'],
        ['W1', '5', 'AI Training', 'Agent Lightning real training pipeline', 'MISSING'],
        ['W1', '-', 'AI MAKER', 'FAKE Voting multi-evaluator consensus', 'MISSING'],
        ['W1', '12', 'Security', 'Rate limit fail-closed on Redis down', 'MISSING'],
        ['W1', '13', 'Security', 'Webhook replay protection (timestamp)', 'MISSING'],
        ['W1', '6', 'Security', 'JWT token blacklist (is_token_revoked)', 'PARTIAL'],
        ['W1', '16', 'Code Quality', 'Shared IP extraction utility', 'MISSING'],
        ['W2', '10', 'AI NLP', 'LLM-based sentiment analysis', 'MISSING'],
        ['W2', '11', 'AI NLP', 'LLM-based PII NER detection', 'MISSING'],
        ['W2', '9', 'AI Voice', 'Voice server LLM intent classification', 'MISSING'],
        ['W2', '8', 'Frontend', '5 Dashboard pages (Billing/KB/Settings/Agents/Monitor)', 'MISSING'],
        ['W2', '3', 'Frontend', 'Socket.io frontend client', 'MISSING'],
        ['W3', '2', 'Security', 'PostgreSQL RLS (93 tables)', 'MISSING'],
        ['W3', '14', 'Database', 'FK mismatch migration fix', 'MISSING'],
        ['W3', '17', 'Infra', 'Nginx non-root USER directive', 'MISSING'],
        ['W3', '21', 'Infra', 'OCSP stapling in docker config', 'PARTIAL'],
        ['W3', '22', 'Infra', 'HSTS includeSubDomains + preload', 'PARTIAL'],
        ['W3', '24', 'Infra', 'TLS cipher suite breadth', 'PARTIAL'],
        ['W3', '23', 'Infra', 'AlertManager diverse receivers', 'MISSING'],
        ['W3', '18', 'Security', 'PRICING_SIGNING_KEY production guard', 'MISSING'],
        ['W3', '19', 'Security', 'ENVIRONMENT enum validation', 'MISSING'],
        ['W3', '7', 'Infra', 'Database backup & recovery', 'MISSING'],
        ['W3', '4', 'Safety', '25 Loophole Solutions engine', 'MISSING'],
        ['W3', '-', 'Security', '12 MEDIUM fixes (M-01/08/13/14/17/19/23/26/28/32/33/35)', 'MISSING'],
        ['W4', '6', 'Security', 'JWT RS256 migration', 'MISSING'],
        ['W4', '-', 'Testing', 'Full test suite (unit + integration + load + failover)', 'NEW'],
        ['W4', '-', 'Deploy', 'Production deployment + SSL + monitoring', 'NEW'],
        ['W4', '-', 'Docs', 'Documentation accuracy update', 'NEW'],
        ['\u2014', '15', 'Database', 'UUID type inconsistency', 'FALSE POSITIVE'],
        ['\u2014', '20', 'Infra', '.env.production.template', 'FALSE POSITIVE'],
        ['\u2014', '-', 'Security', 'M-04 email-validator', 'FALSE POSITIVE'],
        ['\u2014', '-', 'Security', 'M-06 API key 401', 'FALSE POSITIVE'],
        ['\u2014', '-', 'Security', 'M-16 Twilio signature', 'FALSE POSITIVE'],
    ],
    col_widths=[0.4, 0.4, 0.8, 3.5, 1.0]
)

# ══════════════════════════════════════════
# 8. GIT WORKFLOW
# ══════════════════════════════════════════
doc.add_heading('8. Git Workflow and Testing Protocol', level=1)

doc.add_heading('8.1 Strict Git Workflow', level=2)
for item in [
    'At start of every session: git pull origin main to get latest changes',
    'Create feature branch: git checkout -b week-N-task-name (e.g., week-1-clara-rag)',
    'Complete assigned work for the current week',
    'Run the full testing gate for the current week',
    'Fix ALL failing tests before proceeding',
    'Only after ALL tests pass: git push origin week-N-task-name',
    'Create PR, merge to main, pull before starting next week',
]:
    add_bullet(doc, item)

doc.add_heading('8.2 Testing Protocol Per Week', level=2)
for prefix, text in [
    ('1. Unit Tests: ', 'pytest for all new/modified code. Minimum 80% coverage. pytest-cov for reporting.'),
    ('2. Integration Tests: ', 'API endpoints end-to-end, database operations, cross-service communication.'),
    ('3. Security Regression: ', 'Verify no new CRITICAL/HIGH vulnerabilities introduced.'),
    ('4. Performance Baseline: ', 'No API latency or memory degradation from previous week.'),
    ('5. Documentation: ', 'Docstrings on new code. OpenAPI/Swagger updated for new endpoints.'),
]:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_heading('8.3 Parallel Work Handling', level=2)
for item in [
    'Always pull latest main before starting any new work',
    'Resolve merge conflicts immediately, never leave pending',
    'Run full test suite after any merge operation',
    'Never push without passing all tests in the weekly gate',
    'If working on multiple branches, rebase frequently against main',
]:
    add_bullet(doc, item)

# ══════════════════════════════════════════
# 9. EFFORT SUMMARY
# ══════════════════════════════════════════
doc.add_heading('9. Effort Summary', level=1)

add_table(doc,
    ['Week', 'Focus Area', 'Hours', 'Primary Risk'],
    [
        ['Week 1', 'CLARA RAG + Training Pipeline + FAKE Voting + Security', '35-40', 'HyDE prompt quality; training data format'],
        ['Week 2', 'LLM Upgrades (Sentiment/PII/Voice) + 5 Dashboards + Socket.io', '35-40', 'Dashboard API integration; Socket.io auth'],
        ['Week 3', 'RLS + DB Bugs + Nginx + Loopholes + MEDIUM fixes', '35-40', 'RLS migration performance; nginx sync'],
        ['Week 4', 'JWT RS256 + Full Testing + Deployment + Documentation', '30-35', 'RSA key rotation; test coverage gaps'],
        ['TOTAL', '4 Weeks', '135-155', '\u2014'],
    ],
    col_widths=[0.8, 2.8, 0.7, 2.7]
)

add_para(doc, 'This is significantly less than the previous 6-week / 205-235 hour estimate. The reduction comes from eliminating 16 false-positive items that were already built and consolidating the nginx config drift fixes into a single sync operation that resolves 3 items at once.')

# ══════════════════════════════════════════
# 10. SUCCESS METRICS
# ══════════════════════════════════════════
doc.add_heading('10. Success Metrics', level=1)
add_para(doc, 'The roadmap succeeds when the following metrics are achieved at the end of Week 4:')

for m in [
    'CLARA RAG performs HyDE generation and multi-query retrieval (Week 1)',
    'Agent Lightning collects mistakes, triggers training, deploys models (Week 1)',
    'FAKE Voting generates K candidates with multi-evaluator consensus (Week 1)',
    'Rate limiting fails closed when Redis is down (Week 1)',
    'Webhook replay attacks blocked via timestamp validation (Week 1)',
    'JWT token blacklist revokes compromised tokens immediately (Week 1-4)',
    'Sentiment analysis produces LLM-powered structured output (Week 2)',
    'PII detection uses hybrid regex + LLM NER for edge cases (Week 2)',
    'Voice server routes through Smart Router for LLM intent (Week 2)',
    'All 7 dashboard pages display real data from backend APIs (Week 2)',
    'Socket.io provides real-time updates across dashboard (Week 2)',
    'PostgreSQL RLS active on all 93 tenant-scoped tables (Week 3)',
    'Fresh database deployment succeeds with alembic upgrade head (Week 3)',
    'Nginx docker config matches main config (OCSP, HSTS, 8 ciphers) (Week 3)',
    '25 Loophole Solutions detect and counter all defined categories (Week 3)',
    'Automated backups with verified restore capability (Week 3)',
    'All 12 MEDIUM security findings resolved (Week 3)',
    'JWT uses RS256 with key rotation support (Week 4)',
    'Full test suite passes with 80%+ coverage, no regressions (Week 4)',
    '100 concurrent conversations handled stably (Week 4)',
    'SSL Labs A+ test passed (Week 4)',
    'All documentation accurately describes real system capabilities (Week 4)',
]:
    add_bullet(doc, m)

# ── Save ──
output_path = '/home/z/my-project/download/PARWA_Production_Readiness_Roadmap.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')

import os
size = os.path.getsize(output_path)
print(f'File size: {size:,} bytes ({size/1024:.1f} KB)')
